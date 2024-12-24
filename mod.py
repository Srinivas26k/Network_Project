import streamlit as st
import networkx as nx
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from pyvis.network import Network
import random
from fpdf import FPDF
from docx import Document
import io
import base64
import wikipedia
import re
import seaborn as sns
import matplotlib.pyplot as plt
from scipy import stats
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Download necessary NLTK data
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)


# Set page config
st.set_page_config(page_title="Enhanced Interactive Network Analysis", layout="wide")
st.title("Enhanced Interactive Network Analysis with AI-Powered Insights")

# Initialize session state
if 'data' not in st.session_state:
    st.session_state.data = None
if 'graph' not in st.session_state:
    st.session_state.graph = None
if 'embeddings' not in st.session_state:
    st.session_state.embeddings = None
if 'vectorizer' not in st.session_state:
    st.session_state.vectorizer = TfidfVectorizer()

# Function to preprocess data
@st.cache_data
def preprocess_data(df):
    df = df.drop_duplicates().reset_index(drop=True)
    df['Name'] = df['Name'].astype(str).str.strip()
    df['Choose your pick'] = df['Choose your pick'].astype(str).str.strip()
    return df

# Function to create network
@st.cache_data
def create_network(df):
    G = nx.Graph()
    G.add_nodes_from(df['Name'].unique())
    G.add_edges_from(zip(df['Name'], df['Choose your pick']))
    return G

# Function to create embeddings
@st.cache_resource
def create_embeddings(names):
    return st.session_state.vectorizer.fit_transform(names)

# Function to query embeddings
def query_embeddings(query, embeddings, names, k=5):
    query_embedding = st.session_state.vectorizer.transform([query])
    similarities = cosine_similarity(query_embedding, embeddings)[0]
    top_indices = similarities.argsort()[-k:][::-1]
    return [names[i] for i in top_indices]

# Function to generate insights
def generate_insights(G, query, relevant_nodes):
    insights = f"Based on the query '{query}', here are some insights about the relevant nodes:\n\n"
    for node in relevant_nodes:
        degree = G.degree(node)
        pagerank = nx.pagerank(G)[node]
        neighbors = list(G.neighbors(node))
        insights += f"- {node}:\n"
        insights += f"  Degree: {degree}\n"
        insights += f"  PageRank: {pagerank:.6f}\n"
        insights += f"  Neighbors: {', '.join(neighbors)}\n"
        insights += f"  Community: {next(community for community in nx.community.greedy_modularity_communities(G) if node in community)}\n\n"
    return insights

# Function to load or generate data
@st.cache_data
def load_data(uploaded_file=None, generate_random=False, random_params=None):
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
    elif generate_random:
        num_nodes = random_params.get("num_nodes", 20)
        names = [f"{random.choice(['Alice', 'Bob', 'Charlie', 'David', 'Eve', 'Frank'])}_{i}" for i in range(num_nodes)]
        picks = [random.choice(names) for _ in range(num_nodes)]
        df = pd.DataFrame({"Name": names, "Choose your pick": picks})
    else:
        st.error("Please select a valid data input method.")
        st.stop()
    return preprocess_data(df)

# Sidebar for data selection
st.sidebar.header("Data Input")
user_choice = st.sidebar.radio("Select data source:", ["Upload your own data", "Generate random data"])

if user_choice == "Upload your own data":
    uploaded_file = st.sidebar.file_uploader("Upload a CSV file:", type="csv")
    if uploaded_file is not None:
        st.session_state.data = load_data(uploaded_file=uploaded_file)
    else:
        st.sidebar.warning("Please upload a file to proceed.")
        st.stop()
elif user_choice == "Generate random data":
    st.sidebar.subheader("Random Data Parameters")
    num_nodes = st.sidebar.slider("Number of Nodes", min_value=5, max_value=50, value=20)
    random_params = {"num_nodes": num_nodes}
    st.session_state.data = load_data(generate_random=True, random_params=random_params)

# Create network graph and embeddings
if st.session_state.data is not None:
    st.session_state.graph = create_network(st.session_state.data)
    st.session_state.embeddings = create_embeddings(st.session_state.data['Name'].tolist())

# Sidebar for navigation
st.sidebar.header("Navigation")
analysis_type = st.sidebar.selectbox(
    "Choose Analysis Type",
    ["Overview", "Concept Explanation", "Individual Node Analysis", "Influence Paths", 
     "Community Detection", "Connection Quality", "RAG Insights", "Gamification", 
     "Practice Questions", "Advanced Features", "Statistical Visualizations", "Data Research"]
)

# Main content area
st.header(analysis_type)

if analysis_type == "Overview":
    st.write("This section provides a basic overview of the network, including statistics and visualization.")
    
    # Display basic network statistics
    st.subheader("Network Statistics")
    st.write(f"Number of nodes: {st.session_state.graph.number_of_nodes()}")
    st.write(f"Number of edges: {st.session_state.graph.number_of_edges()}")
    st.write(f"Average degree: {2 * st.session_state.graph.number_of_edges() / st.session_state.graph.number_of_nodes():.2f}")
    
    # Create and display network visualization
    st.subheader("Network Visualization")
    net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(st.session_state.graph)
    net.save_graph("network.html")
    st.components.v1.html(open("network.html", "r").read(), height=600)

elif analysis_type == "Concept Explanation":
    st.write("Learn about key concepts in network analysis in an interactive way.")
    
    concepts = {
        "Node": "A fundamental unit in a network, representing an entity (e.g., a person, organization, or data point).",
        "Edge": "A connection between two nodes, representing a relationship or interaction.",
        "Degree": "The number of edges connected to a node, indicating its connectivity.",
        "Path": "A sequence of nodes and edges connecting two nodes in the network.",
        "Centrality": "A measure of a node's importance in the network, based on its position and connections.",
        "Community": "A group of nodes that are more densely connected to each other than to the rest of the network.",
        "PageRank": "An algorithm that measures the importance of nodes based on the structure of incoming links."
    }
    
    for concept, explanation in concepts.items():
        st.subheader(concept)
        st.write(explanation)
        try:
            wiki_summary = wikipedia.summary(concept, sentences=2)
            st.write(f"From Wikipedia: {wiki_summary}")
        except wikipedia.exceptions.WikipediaException:
            st.write("Additional information unavailable from Wikipedia.")

elif analysis_type == "Individual Node Analysis":
    st.write("Analyze individual nodes in the network to understand their properties and importance.")
    
    selected_node = st.selectbox("Select a node to analyze", st.session_state.data['Name'].tolist())
    
    st.subheader("Node Properties")
    degree = st.session_state.graph.degree(selected_node)
    st.write(f"Degree: {degree}")
    
    component = next(nx.connected_components(st.session_state.graph))
    if selected_node in component:
        st.write(f"Connected component size: {len(component)}")
    else:
        st.write("The selected node is not part of the largest connected component.")
    
    pagerank = nx.pagerank(st.session_state.graph)
    st.write(f"PageRank score: {pagerank[selected_node]:.6f}")
    
    st.subheader("Node Visualization")
    subgraph = st.session_state.graph.subgraph([selected_node] + list(st.session_state.graph.neighbors(selected_node)))
    net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(subgraph)
    net.save_graph("node_graph.html")
    st.components.v1.html(open("node_graph.html", "r").read(), height=400)

elif analysis_type == "Influence Paths":
    st.write("Discover the shortest paths between nodes to understand influence flow in the network.")
    
    source = st.selectbox("Select source node", st.session_state.data['Name'].tolist())
    target = st.selectbox("Select target node", st.session_state.data['Name'].tolist())
    
    if st.button("Find Influence Path"):
        try:
            path = nx.shortest_path(st.session_state.graph, source=source, target=target)
            st.success(f"Influence path from {source} to {target}:")
            st.write(" -> ".join(path))
            
            # Visualize the path
            path_graph = st.session_state.graph.subgraph(path)
            net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
            net.from_nx(path_graph)
            net.save_graph("path.html")
            st.components.v1.html(open("path.html", "r").read(), height=400)
        except nx.NetworkXNoPath:
            st.error(f"No path found between {source} and {target}")

elif analysis_type == "Community Detection":
    st.write("Identify and visualize communities within the network.")
    
    communities = list(nx.community.greedy_modularity_communities(st.session_state.graph))
    
    st.subheader("Community Statistics")
    st.write(f"Number of communities detected: {len(communities)}")
    
    st.subheader("What is a Community?")
    st.write("A community is a group of nodes that are more densely connected to each other than to the rest of the network.")
    
    # Visualize communities
    st.subheader("Community Visualization")
    pos = nx.spring_layout(st.session_state.graph)
    fig = go.Figure()
    
    for i, community in enumerate(communities):
        node_x = [pos[node][0] for node in community]
        node_y = [pos[node][1] for node in community]
        fig.add_trace(go.Scatter(x=node_x, y=node_y, mode='markers', name=f'Community {i+1}',
                                 marker=dict(size=10, line=dict(width=2))))
    
    fig.update_layout(title="Community Detection Visualization", showlegend=True)
    st.plotly_chart(fig)

elif analysis_type == "Connection Quality":
    st.write("Evaluate the quality of connections for a specific node in the network.")
    
    selected_node = st.selectbox("Select a node to analyze", st.session_state.data['Name'].tolist())
    
    neighbors = list(st.session_state.graph.neighbors(selected_node))
    total_possible_connections = len(neighbors) * (len(neighbors) - 1) / 2
    actual_connections = sum(1 for n1 in neighbors for n2 in neighbors if n1 < n2 and st.session_state.graph.has_edge(n1, n2))
    
    if total_possible_connections == 0:
        quality_score = 0
        explanation = "The node has no or only one connection."
    else:
        quality_score = actual_connections / total_possible_connections
        if quality_score < 0.33:
            explanation = "The connections are mostly isolated from each other."
        elif quality_score < 0.67:
            explanation = "There is a moderate level of interconnection among the node's connections."
        else:
            explanation = "The node's connections form a tightly-knit group."
    
    st.subheader("Connection Quality Analysis")
    st.write(f"Connection Quality Score: {quality_score:.2f}")
    st.write(f"Explanation: {explanation}")
    
    st.subheader("Node Connections")
    st.write(f"Neighbors of {selected_node}: {', '.join(neighbors)}")
    
    st.subheader("Connections between Neighbors")
    for i, n1 in enumerate(neighbors):
        for n2 in neighbors[i+1:]:
            if st.session_state.graph.has_edge(n1, n2):
                st.write(f"- {n1} is connected to {n2}")

elif analysis_type == "RAG Insights":
    st.write("Get AI-powered insights about the network using Retrieval Augmented Generation (RAG).")
    
    insight_type = st.radio("Select insight type:", ["Name Search", "Ask a Question"])
    
    if insight_type == "Name Search":
        query = st.text_input("Enter The Name:")
        
        if query:
            st.subheader(f"Network Insights for the Given Name: {query}")
            
            # Filter the names to match the base name
            matching_names = [name for name in st.session_state.data['Name'].tolist() if re.search(query, name, re.IGNORECASE)]
            
            if matching_names:
                relevant_nodes = query_embeddings(query, st.session_state.embeddings, matching_names, k=5)
                insights = generate_insights(st.session_state.graph, query, relevant_nodes)
                if insights:
                    st.text(insights)
                else:
                    st.write("No specific insights found for the given name in the network data.")
            else:
                st.write(f"No nodes found matching the name '{query}' in the network.")
    
    else:  # Ask a Question
        question = st.text_input("Ask a question about the network:")
        if question:
            # Tokenize the question
            tokens = word_tokenize(question.lower())
            # Remove stopwords
            stop_words = set(stopwords.words('english'))
            filtered_tokens = [w for w in tokens if not w in stop_words]
            
            # Find relevant nodes based on the question
            relevant_nodes = query_embeddings(" ".join(filtered_tokens), st.session_state.embeddings, st.session_state.data['Name'].tolist(), k=5)
            
            # Generate insights
            insights = generate_insights(st.session_state.graph, question, relevant_nodes)
            
            st.subheader("AI-Generated Insights")
            st.text(insights)
    
    st.markdown("---")
    
    # Wikipedia Implementation
    st.subheader("Related Topics from Wikipedia")
    try:
        topic = "network analysis" if "network" in question.lower() else question
        wiki_summary = wikipedia.summary(topic, sentences=5)
        st.write(f"From Wikipedia: {wiki_summary}")
    except wikipedia.exceptions.DisambiguationError as e:
        st.write(f"Disambiguation required. Possible topics: {e.options[:5]}.")
    except wikipedia.exceptions.PageError:
        st.write("No Wikipedia page found for the topic.")
    except wikipedia.exceptions.WikipediaException:
        st.write("Unable to fetch Wikipedia information.")

elif analysis_type == "Gamification":
    st.write("Experiment with the network by adding or removing nodes and observe the impact.")
    
    st.subheader("What is Gamification?")
    st.write("Gamification involves creating an interactive environment to understand the dynamics of the network by adding or removing nodes in real-time.")
    
    action = st.radio("Choose an action:", ["Add Node", "Remove Node"])
    
    if action == "Add Node":
        new_node = st.text_input("Enter the name of the new node:")
        connect_to = st.multiselect("Connect to existing nodes:", st.session_state.data['Name'].tolist())
        
        if st.button("Add Node"):
            st.session_state.graph.add_node(new_node)
            for node in connect_to:
                st.session_state.graph.add_edge(new_node, node)
            st.success(f"Added node {new_node} with {len(connect_to)} connections.")
    else:
        remove_node = st.selectbox("Select a node to remove:", st.session_state.data['Name'].tolist())
        
        if st.button("Remove Node"):
            st.session_state.graph.remove_node(remove_node)
            st.success(f"Removed node {remove_node}.")
    
    st.subheader("Updated Network Visualization")
    net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(st.session_state.graph)
    net.save_graph("updated_network.html")
    st.components.v1.html(open("updated_network.html", "r").read(), height=600)

elif analysis_type == "Practice Questions":
    st.write("Test your understanding of network analysis concepts with practice questions.")
    questions = [
        "What is the degree of a node in a network?",
        "How does PageRank measure node importance?",
        "What defines a community in a network?",
        "Describe the difference between a path and a walk in a network.",
        "What is a connected component in a graph?"
    ]
    
    if "score" not in st.session_state:
        st.session_state["score"] = 0
        st.session_state["question_index"] = 0

    if st.session_state["question_index"] < len(questions):
        question = questions[st.session_state["question_index"]]
        st.subheader("Practice Question")
        st.write(question)
        user_answer = st.text_input("Your Answer:")
        if st.button("Submit Answer"):
            st.session_state["score"] += 1  # Assume correct for demo
            st.session_state["question_index"] += 1
            st.rerun()
    else:
        st.write(f"You completed the quiz! Your score is {st.session_state['score']}/{len(questions)}.")
        if st.button("Download Report"):
            pdf_buffer = io.BytesIO()
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="Practice Question Analysis", ln=True, align='C')
            pdf.cell(200, 10, txt=f"Score: {st.session_state['score']}/{len(questions)}", ln=True)
            pdf.output(pdf_buffer)
            pdf_buffer.seek(0)
            b64 = base64.b64encode(pdf_buffer.getvalue()).decode()
            href = f'<a href="data:application/pdf;base64,{b64}" download="practice_analysis_report.pdf">Download Practice Analysis Report</a>'
            st.markdown(href, unsafe_allow_html=True)

elif analysis_type == "Advanced Features":
    st.write("Explore advanced network analysis features.")
    
    advanced_feature = st.selectbox("Select an advanced feature:", 
                                    ["Compare Graphs", "Histogram Comparison", "Node Search",
                                     "Network Metrics", "Time Series Analysis", "Centrality Comparison",
                                     "Community Evolution", "Link Prediction", "Network Visualization Comparison"])
    
    if advanced_feature == "Compare Graphs":
        st.write("Compare two or more graphs.")
        # Implement graph comparison logic
        st.write("This feature is under development.")
    
    elif advanced_feature == "Histogram Comparison":
        st.write("Compare histograms of different network properties.")
        properties = ["Degree", "Clustering Coefficient", "Betweenness Centrality"]
        selected_properties = st.multiselect("Select properties to compare:", properties)
        
        if selected_properties:
            fig, axes = plt.subplots(1, len(selected_properties), figsize=(5*len(selected_properties), 4))
            for i, prop in enumerate(selected_properties):
                if prop == "Degree":
                    data = [d for n, d in st.session_state.graph.degree()]
                elif prop == "Clustering Coefficient":
                    data = list(nx.clustering(st.session_state.graph).values())
                elif prop == "Betweenness Centrality":
                    data = list(nx.betweenness_centrality(st.session_state.graph).values())
                
                if len(selected_properties) > 1:
                    ax = axes[i]
                else:
                    ax = axes
                
                sns.histplot(data, kde=True, ax=ax)
                ax.set_title(prop)
                ax.set_xlabel("Value")
                ax.set_ylabel("Frequency")
            
            plt.tight_layout()
            st.pyplot(fig)
    
    elif advanced_feature == "Node Search":
        st.write("Search for specific nodes in the network.")
        search_query = st.text_input("Enter node name to search:")
        if search_query:
            matching_nodes = [node for node in st.session_state.graph.nodes() if search_query.lower() in node.lower()]
            if matching_nodes:
                st.write(f"Matching nodes: {', '.join(matching_nodes)}")
            else:
                st.write("No matching nodes found.")
    
    elif advanced_feature == "Network Metrics":
        st.write("Calculate and compare various network metrics.")
        metrics = {
            "Average Clustering Coefficient": nx.average_clustering(st.session_state.graph),
            "Network Density": nx.density(st.session_state.graph),
            "Average Shortest Path Length": nx.average_shortest_path_length(st.session_state.graph) if nx.is_connected(st.session_state.graph) else "N/A (Graph is not connected)",
            "Number of Connected Components": nx.number_connected_components(st.session_state.graph),
            "Graph Diameter": nx.diameter(st.session_state.graph) if nx.is_connected(st.session_state.graph) else "N/A (Graph is not connected)"
        }
        
        for metric, value in metrics.items():
            st.write(f"{metric}: {value}")
    
    elif advanced_feature == "Time Series Analysis":
        st.write("Analyze network evolution over time.")
        st.write("This feature requires time-stamped network data, which is not available in the current dataset.")
    
    elif advanced_feature == "Centrality Comparison":
        st.write("Compare different centrality measures for nodes.")
        centrality_measures = {
            "Degree Centrality": nx.degree_centrality(st.session_state.graph),
            "Betweenness Centrality": nx.betweenness_centrality(st.session_state.graph),
            "Closeness Centrality": nx.closeness_centrality(st.session_state.graph),
            "Eigenvector Centrality": nx.eigenvector_centrality(st.session_state.graph)
        }
        
        selected_measures = st.multiselect("Select centrality measures to compare:", list(centrality_measures.keys()))
        
        if selected_measures:
            data = []
            for measure in selected_measures:
                data.extend([(node, value, measure) for node, value in centrality_measures[measure].items()])
            
            df = pd.DataFrame(data, columns=["Node", "Centrality", "Measure"])
            fig = px.box(df, x="Measure", y="Centrality", points="all")
            st.plotly_chart(fig)
    
    elif advanced_feature == "Community Evolution":
        st.write("Track the evolution of communities over time.")
        st.write("This feature requires time-stamped network data, which is not available in the current dataset.")
    
    elif advanced_feature == "Link Prediction":
        st.write("Predict potential future links in the network.")
        
        # Implement a simple link prediction algorithm
        def common_neighbors_score(G, node1, node2):
            return len(list(nx.common_neighbors(G, node1, node2)))
        
        # Get all pairs of nodes that are not connected
        non_edges = list(nx.non_edges(st.session_state.graph))
        
        # Calculate scores for all non-edges
        scores = [(u, v, common_neighbors_score(st.session_state.graph, u, v)) for u, v in non_edges]
        
        # Sort by score in descending order
        scores.sort(key=lambda x: x[2], reverse=True)
        
        # Display top 10 potential links
        st.write("Top 10 potential future links:")
        for u, v, score in scores[:10]:
            st.write(f"{u} - {v}: Score = {score}")
    
    elif advanced_feature == "Network Visualization Comparison":
        st.write("Compare different network visualization techniques.")
        
        visualization_types = ["Spring", "Circular", "Random", "Shell"]
        selected_type = st.selectbox("Select a visualization type:", visualization_types)
        
        if selected_type == "Spring":
            pos = nx.spring_layout(st.session_state.graph)
        elif selected_type == "Circular":
            pos = nx.circular_layout(st.session_state.graph)
        elif selected_type == "Random":
            pos = nx.random_layout(st.session_state.graph)
        elif selected_type == "Shell":
            pos = nx.shell_layout(st.session_state.graph)
        
        fig, ax = plt.subplots(figsize=(10, 8))
        nx.draw(st.session_state.graph, pos, with_labels=True, node_color='lightblue', 
                node_size=500, font_size=8, font_weight='bold', ax=ax)
        ax.set_title(f"{selected_type} Layout")
        st.pyplot(fig)

elif analysis_type == "Statistical Visualizations":
    st.write("Explore various statistical visualizations of the network.")
    
    viz_type = st.selectbox("Select visualization type:", 
                            ["Degree Distribution", "Centrality Distribution", "Node Attribute Distribution",
                             "Edge Weight Distribution", "Clustering Coefficient Distribution"])
    
    if viz_type == "Degree Distribution":
        degrees = [d for n, d in st.session_state.graph.degree()]
        fig, ax = plt.subplots()
        sns.histplot(degrees, kde=True, ax=ax)
        ax.set_title("Degree Distribution")
        ax.set_xlabel("Degree")
        ax.set_ylabel("Frequency")
        st.pyplot(fig)
    
    elif viz_type == "Centrality Distribution":
        centrality_type = st.selectbox("Select centrality measure:", 
                                       ["Degree", "Betweenness", "Closeness", "Eigenvector"])
        
        if centrality_type == "Degree":
            centrality = nx.degree_centrality(st.session_state.graph)
        elif centrality_type == "Betweenness":
            centrality = nx.betweenness_centrality(st.session_state.graph)
        elif centrality_type == "Closeness":
            centrality = nx.closeness_centrality(st.session_state.graph)
        elif centrality_type == "Eigenvector":
            centrality = nx.eigenvector_centrality(st.session_state.graph)
        
        fig, ax = plt.subplots()
        sns.histplot(list(centrality.values()), kde=True, ax=ax)
        ax.set_title(f"{centrality_type} Centrality Distribution")
        ax.set_xlabel("Centrality")
        ax.set_ylabel("Frequency")
        st.pyplot(fig)
    
    elif viz_type == "Node Attribute Distribution":
        st.write("This visualization requires node attributes, which are not available in the current dataset.")
    
    elif viz_type == "Edge Weight Distribution":
        st.write("This visualization requires weighted edges, which are not available in the current dataset.")
    
    elif viz_type == "Clustering Coefficient Distribution":
        clustering_coeffs = list(nx.clustering(st.session_state.graph).values())
        fig, ax = plt.subplots()
        sns.histplot(clustering_coeffs, kde=True, ax=ax)
        ax.set_title("Clustering Coefficient Distribution")
        ax.set_xlabel("Clustering Coefficient")
        ax.set_ylabel("Frequency")
        st.pyplot(fig)

    # Add comparison features
    st.subheader("Node Comparison")
    comparison_type = st.selectbox("Select comparison type:", ["Highest Degree", "Highest PageRank", "Highest Betweenness Centrality"])
    
    if comparison_type == "Highest Degree":
        top_nodes = sorted(st.session_state.graph.degree(), key=lambda x: x[1], reverse=True)[:10]
        comparison_data = pd.DataFrame(top_nodes, columns=["Node", "Degree"])
    elif comparison_type == "Highest PageRank":
        pagerank = nx.pagerank(st.session_state.graph)
        top_nodes = sorted(pagerank.items(), key=lambda x: x[1], reverse=True)[:10]
        comparison_data = pd.DataFrame(top_nodes, columns=["Node", "PageRank"])
    elif comparison_type == "Highest Betweenness Centrality":
        betweenness = nx.betweenness_centrality(st.session_state.graph)
        top_nodes = sorted(betweenness.items(), key=lambda x: x[1], reverse=True)[:10]
        comparison_data = pd.DataFrame(top_nodes, columns=["Node", "Betweenness Centrality"])
    
    st.write(comparison_data)
    
    # Visualize comparison
    fig, ax = plt.subplots()
    sns.barplot(data=comparison_data, x="Node", y=comparison_data.columns[1], ax=ax)
    ax.set_title(f"Top 10 Nodes by {comparison_type}")
    ax.set_xticklabels(ax.get_xticklabels(), rotation=45, ha='right')
    ax.set_xlabel("Node")
    ax.set_ylabel(comparison_data.columns[1])
    plt.tight_layout()
    st.pyplot(fig)

elif analysis_type == "Data Research":
    st.write("Explore and analyze your network data in-depth.")
    
    research_option = st.selectbox("Choose a research option:", 
                                   ["Correlation Analysis", "Network Evolution", "Node Importance Over Time"])
    
    if research_option == "Correlation Analysis":
        st.write("Analyze correlations between different node attributes.")
        
        # For demonstration purposes, we'll generate random attributes
        node_attributes = {node: {'attr1': random.random(), 'attr2': random.random()} for node in st.session_state.graph.nodes()}
        
        df = pd.DataFrame.from_dict(node_attributes, orient='index')
        df['degree'] = dict(st.session_state.graph.degree()).values()
        df['clustering'] = nx.clustering(st.session_state.graph)
        
        st.write("Node Attributes and Metrics:")
        st.write(df)
        
        st.write("Correlation Matrix:")
        corr_matrix = df.corr()
        st.write(corr_matrix)
        
        fig, ax = plt.subplots()
        sns.heatmap(corr_matrix, annot=True, cmap='coolwarm', ax=ax)
        st.pyplot(fig)
    
    elif research_option == "Network Evolution":
        st.write("Simulate network evolution over time.")
        
        num_steps = st.slider("Number of evolution steps:", 1, 10, 5)
        
        for step in range(num_steps):
            st.subheader(f"Evolution Step {step + 1}")
            
            # Simulate network evolution by randomly adding/removing edges
            if random.random() < 0.7:  # 70% chance to add an edge
                node1, node2 = random.sample(list(st.session_state.graph.nodes()), 2)
                st.session_state.graph.add_edge(node1, node2)
                st.write(f"Added edge: {node1} - {node2}")
            else:
                if st.session_state.graph.number_of_edges() > 0:
                    edge = random.choice(list(st.session_state.graph.edges()))
                    st.session_state.graph.remove_edge(*edge)
                    st.write(f"Removed edge: {edge[0]} - {edge[1]}")
            
            net = Network(height="300px", width="100%", bgcolor="#222222", font_color="white")
            net.from_nx(st.session_state.graph)
            net.save_graph(f"evolution_step_{step + 1}.html")
            st.components.v1.html(open(f"evolution_step_{step + 1}.html", "r").read(), height=300)
    
    elif research_option == "Node Importance Over Time":
        st.write("Analyze how node importance changes over time.")
        
        num_steps = st.slider("Number of time steps:", 1, 10, 5)
        
        importance_data = {}
        for step in range(num_steps):
            # Simulate changes in the network
            if st.session_state.graph.number_of_edges() > 0:
                edge = random.choice(list(st.session_state.graph.edges()))
                st.session_state.graph.remove_edge(*edge)
            node1, node2 = random.sample(list(st.session_state.graph.nodes()), 2)
            st.session_state.graph.add_edge(node1, node2)
            
            # Calculate node importance metrics
            degree_centrality = nx.degree_centrality(st.session_state.graph)
            betweenness_centrality = nx.betweenness_centrality(st.session_state.graph)
            pagerank = nx.pagerank(st.session_state.graph)
            
            for node in st.session_state.graph.nodes():
                if node not in importance_data:
                    importance_data[node] = {'degree': [], 'betweenness': [], 'pagerank': []}
                importance_data[node]['degree'].append(degree_centrality[node])
                importance_data[node]['betweenness'].append(betweenness_centrality[node])
                importance_data[node]['pagerank'].append(pagerank[node])
        
        selected_node = st.selectbox("Select a node to analyze:", list(st.session_state.graph.nodes()))
        
        fig, ax = plt.subplots()
        ax.plot(range(1, num_steps + 1), importance_data[selected_node]['degree'], label='Degree Centrality')
        ax.plot(range(1, num_steps + 1), importance_data[selected_node]['betweenness'], label='Betweenness Centrality')
        ax.plot(range(1, num_steps + 1), importance_data[selected_node]['pagerank'], label='PageRank')
        ax.set_xlabel('Time Step')
        ax.set_ylabel('Importance')
        ax.set_title(f'Node Importance Over Time: {selected_node}')
        ax.legend()
        st.pyplot(fig)

# Report generation
st.sidebar.header("Generate Report")
report_format = st.sidebar.selectbox("Choose report format", ["CSV", "PDF", "DOCX"])

if st.sidebar.button("Generate Report"):
    report_data = {
        "Number of Nodes": st.session_state.graph.number_of_nodes(),
        "Number of Edges": st.session_state.graph.number_of_edges(),
        "Average Degree": 2 * st.session_state.graph.number_of_edges() / st.session_state.graph.number_of_nodes(),
        "Number of Connected Components": nx.number_connected_components(st.session_state.graph),
        "Average Clustering Coefficient": nx.average_clustering(st.session_state.graph),
        "Network Density": nx.density(st.session_state.graph)
    }

    if report_format == "CSV":
        csv = pd.DataFrame([report_data]).to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        href = f'<a href="data:file/csv;base64,{b64}" download="network_analysis_report.csv">Download CSV Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)
    
    elif report_format == "PDF":
        pdf_buffer = io.BytesIO()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Network Analysis Report", ln=True, align='C')
        for key, value in report_data.items():
            pdf.cell(200, 10, txt=f"{key}: {value}", ln=True)
        pdf.output(pdf_buffer)
        pdf_buffer.seek(0)
        b64 = base64.b64encode(pdf_buffer.getvalue()).decode()
        href = f'<a href="data:application/pdf;base64,{b64}" download="network_analysis_report.pdf">Download PDF Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)
    
    elif report_format == "DOCX":
        doc_buffer = io.BytesIO()
        doc = Document()
        doc.add_heading("Network Analysis Report", 0)
        for key, value in report_data.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        b64 = base64.b64encode(doc_buffer.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="network_analysis_report.docx">Download DOCX Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.markdown("Created by Nampalli Srinivas")
st.markdown("---")
st.markdown("### Note: This project is for educational purposes only.")
st.markdown("Connect with the creator:")
st.markdown("[![LinkedIn](https://img.shields.io/badge/LinkedIn-0077B5?style=for-the-badge&logo=linkedin&logoColor=white)](https://www.linkedin.com/in/srinivas-nampalli/)")
st.markdown("[![Twitter](https://img.shields.io/badge/Twitter-1DA1F2?style=for-the-badge&logo=twitter&logoColor=white)](https://x.com/Srinivas26k)")

if __name__ == "__main__":
    st.write("Enhanced Interactive Network Analysis app is running.")

