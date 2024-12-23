import streamlit as st
import networkx as nx
import pandas as pd
import plotly.graph_objects as go
from pyvis.network import Network
import faiss
import numpy as np
import random
from fpdf import FPDF
from docx import Document
import io
import base64
from utils import preprocess_data, create_network, create_embeddings, query_faiss, generate_report, calculate_pagerank, analyze_connection_quality, generate_insights
import wikipedia

# Set up FAISS
dimension = 1  # Dimension of the simple embeddings
index = faiss.IndexFlatL2(dimension)

st.set_page_config(page_title="Interactive Network Analysis", layout="wide")
st.title("Interactive Network Analysis with AI-Powered Insights")

# Feature: Allow user to upload or generate random data
@st.cache_data
def load_data(uploaded_file=None, generate_random=False, random_params=None):
    if uploaded_file is not None:
        df = pd.read_csv(uploaded_file)
    elif generate_random:
        # Generate random data with flexible parameters
        if random_params:
            num_nodes = random_params.get("num_nodes", 20)
            names = [f"{random.choice(['Alice', 'Bob', 'Charlie', 'David', 'Eve', 'Frank'])}_{i}" for i in range(num_nodes)]
            picks = [random.choice(names) for _ in range(num_nodes)]
            df = pd.DataFrame({"Name": names, "Choose your pick": picks})
        else:
            # Default random generation
            names = [f"User_{i}" for i in range(20)]
            picks = np.random.choice(names, size=20)
            df = pd.DataFrame({"Name": names, "Choose your pick": picks})
    else:
        st.error("Please select a valid data input method.")
        st.stop()
    return preprocess_data(df)

# Sidebar for data selection
st.sidebar.header("Data Input")
user_choice = st.sidebar.radio("Select data source:", ["Upload your own data", "Generate random data"])

if user_choice == "Upload your own data":
    uploaded_file = st.sidebar.file_uploader("Upload a CSV file:", type="csv")
    if uploaded_file is not None:
        df = load_data(uploaded_file=uploaded_file)
    else:
        st.sidebar.warning("Please upload a file to proceed.")
        st.stop()
elif user_choice == "Generate random data":
    st.sidebar.subheader("Random Data Parameters")
    num_nodes = st.sidebar.slider("Number of Nodes", min_value=5, max_value=50, value=20)
    random_params = {"num_nodes": num_nodes}
    df = load_data(generate_random=True, random_params=random_params)

# Create network graph
G = create_network(df)

# Create embeddings and add to FAISS index
embeddings = create_embeddings(df['Name'].tolist())
index.add(embeddings)

# Sidebar for navigation
st.sidebar.header("Navigation")
analysis_type = st.sidebar.selectbox(
    "Choose Analysis Type",
    ["Overview", "Concept Explanation", "Individual Node Analysis", "Influence Paths", "Community Detection", "Connection Quality", "RAG Insights", "Gamification", "Practice Questions"]
)

# Main content area
st.header(analysis_type)

if analysis_type == "Overview":
    st.write("This section provides a basic overview of the network, including statistics and visualization.")
    
    # Display basic network statistics
    st.subheader("Network Statistics")
    st.write(f"Number of nodes: {G.number_of_nodes()}")
    st.write(f"Number of edges: {G.number_of_edges()}")
    st.write(f"Average degree: {2 * G.number_of_edges() / G.number_of_nodes():.2f}")
    
    # Create and display network visualization
    st.subheader("Network Visualization")
    net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(G)
    net.save_graph("network.html")
    st.components.v1.html(open("network.html", "r").read(), height=600)

elif analysis_type == "Concept Explanation":
    st.write("Learn about key concepts in network analysis in an interactive way.")
    
    concepts = {
        "Node": "A fundamental unit in a network, representing an entity (e.g., a person, organization, or data point).",
        "Edge": "A connection between two nodes, representing a relationship or interaction.",
        "Degree": "The number of edges connected to a node, indicating its connectivity.",
        "Path": "A sequence of nodes and edges connecting two nodes in the network.",
        "Centrality": "A measure of a node's importance in the network, based on its position and connections.",
        "Community": "A group of nodes that are more densely connected to each other than to the rest of the network.",
        "PageRank": "An algorithm that measures the importance of nodes based on the structure of incoming links."
    }
    
    for concept, explanation in concepts.items():
        st.subheader(concept)
        st.write(explanation)
        try:
            wiki_summary = wikipedia.summary(concept, sentences=2)
            st.write(f"From Wikipedia: {wiki_summary}")
        except wikipedia.exceptions.WikipediaException:
            st.write("Additional information unavailable from Wikipedia.")

elif analysis_type == "RAG Insights":
    st.write("Get AI-powered insights about the network using Retrieval Augmented Generation (RAG).")
    
    query = st.text_input("Ask a question about the network:")
    
    if query:
        st.subheader("Network Insights")
        relevant_nodes = query_faiss(index, query, df['Name'].tolist(), k=5)
        insights = generate_insights(G, query, relevant_nodes)
        if insights:
            st.write(insights)
        else:
            st.write("No specific insights found for your query in the network data.")
        st.markdown("---")
        st.subheader("Related Topics from Wikipedia")
        try:
            topic = "network" if "network" in query.lower() else query
            wiki_summary = wikipedia.summary(topic, sentences=10)
            st.write(f"From Wikipedia: {wiki_summary}")
        except wikipedia.exceptions.DisambiguationError as e:
            st.write(f"Disambiguation required. Possible topics: {e.options[:5]}.")
        except wikipedia.exceptions.PageError:
            st.write("No Wikipedia page found for the topic.")
        except wikipedia.exceptions.WikipediaException:
            st.write("Unable to fetch Wikipedia information.")

elif analysis_type == "Individual Node Analysis":
    st.write("Analyze individual nodes in the network to understand their properties and importance.")
    
    selected_node = st.selectbox("Select a node to analyze", df['Name'].tolist())
    
    st.subheader("Node Properties")
    degree = G.degree(selected_node)
    st.write(f"Degree: {degree}")
    
    component = next(nx.connected_components(G))
    if selected_node in component:
        st.write(f"Connected component size: {len(component)}")
    else:
        st.write("The selected node is not part of the largest connected component.")
    
    pagerank = calculate_pagerank(G)
    st.write(f"PageRank score: {pagerank[selected_node]:.6f}")
    
    st.subheader("Node Visualization")
    subgraph = G.subgraph([selected_node] + list(G.neighbors(selected_node)))
    net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(subgraph)
    net.save_graph("node_graph.html")
    st.components.v1.html(open("node_graph.html", "r").read(), height=400)

elif analysis_type == "Influence Paths":
    st.write("Discover the shortest paths between nodes to understand influence flow in the network.")
    
    source = st.selectbox("Select source node", df['Name'].tolist())
    target = st.selectbox("Select target node", df['Name'].tolist())
    
    if st.button("Find Influence Path"):
        try:
            path = nx.shortest_path(G, source=source, target=target)
            st.success(f"Influence path from {source} to {target}:")
            st.write(" -> ".join(path))
            
            # Visualize the path
            path_graph = G.subgraph(path)
            net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
            net.from_nx(path_graph)
            net.save_graph("path.html")
            st.components.v1.html(open("path.html", "r").read(), height=400)
        except nx.NetworkXNoPath:
            st.error(f"No path found between {source} and {target}")

elif analysis_type == "Community Detection":
    st.write("Identify and visualize communities within the network.")
    
    communities = nx.community.louvain_communities(G)
    
    st.subheader("Community Statistics")
    st.write(f"Number of communities detected: {len(communities)}")
    
    st.subheader("What is a Community?")
    st.write("A community is a group of nodes that are more densely connected to each other than to the rest of the network.")
    
    # Visualize communities
    st.subheader("Community Visualization")
    pos = nx.spring_layout(G)
    fig = go.Figure()
    
    for i, community in enumerate(communities):
        node_x = [pos[node][0] for node in community]
        node_y = [pos[node][1] for node in community]
        fig.add_trace(go.Scatter(x=node_x, y=node_y, mode='markers', name=f'Community {i+1}',
                                 marker=dict(size=10, line=dict(width=2))))
    
    fig.update_layout(title="Community Detection Visualization", showlegend=True)
    st.plotly_chart(fig)

elif analysis_type == "Connection Quality":
    st.write("Evaluate the quality of connections for a specific node in the network.")
    
    selected_node = st.selectbox("Select a node to analyze", df['Name'].tolist())
    quality_score, explanation = analyze_connection_quality(G, selected_node)
    
    st.subheader("Connection Quality Analysis")
    st.write(f"Connection Quality Score: {quality_score:.2f}")
    st.write(f"Explanation: {explanation}")
    
    st.subheader("Node Connections")
    neighbors = list(G.neighbors(selected_node))
    st.write(f"Neighbors of {selected_node}: {', '.join(neighbors)}")
    
    st.subheader("Connections between Neighbors")
    for i, n1 in enumerate(neighbors):
        for n2 in neighbors[i+1:]:
            if G.has_edge(n1, n2):
                st.write(f"- {n1} is connected to {n2}")

elif analysis_type == "Gamification":
    st.write("Experiment with the network by adding or removing nodes and observe the impact.")
    
    st.subheader("What is Gamification?")
    st.write("Gamification involves creating an interactive environment to understand the dynamics of the network by adding or removing nodes in real-time.")
    
    action = st.radio("Choose an action:", ["Add Node", "Remove Node"])
    
    if action == "Add Node":
        new_node = st.text_input("Enter the name of the new node:")
        connect_to = st.multiselect("Connect to existing nodes:", df['Name'].tolist())
        
        if st.button("Add Node"):
            G.add_node(new_node)
            for node in connect_to:
                G.add_edge(new_node, node)
            st.success(f"Added node {new_node} with {len(connect_to)} connections.")
    else:
        remove_node = st.selectbox("Select a node to remove:", df['Name'].tolist())
        
        if st.button("Remove Node"):
            G.remove_node(remove_node)
            st.success(f"Removed node {remove_node}.")
    
    st.subheader("Updated Network Visualization")
    net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(G)
    net.save_graph("updated_network.html")
    st.components.v1.html(open("updated_network.html", "r").read(), height=600)

elif analysis_type == "Practice Questions":
    st.write("Test your understanding of network analysis concepts with practice questions.")
    questions = [
        "What is the degree of a node in a network?",
        "How does PageRank measure node importance?",
        "What defines a community in a network?",
        "Describe the difference between a path and a walk in a network.",
        "What is a connected component in a graph?"
    ]
    
    if "score" not in st.session_state:
        st.session_state["score"] = 0
        st.session_state["question_index"] = 0

    if st.session_state["question_index"] < len(questions):
        question = questions[st.session_state["question_index"]]
        st.subheader("Practice Question")
        st.write(question)
        user_answer = st.text_input("Your Answer:")
        if st.button("Submit Answer"):
            st.session_state["score"] += 1  # Assume correct for demo
            st.session_state["question_index"] += 1
            st.experimental_rerun()  # Properly trigger a rerun
    else:
        st.write(f"You completed the quiz! Your score is {st.session_state['score']}/{len(questions)}.")
        if st.button("Download Report"):
            pdf_buffer = io.BytesIO()
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="Practice Question Analysis", ln=True, align='C')
            pdf.cell(200, 10, txt=f"Score: {st.session_state['score']}/{len(questions)}", ln=True)
            pdf.output(pdf_buffer)
            pdf_buffer.seek(0)
            b64 = base64.b64encode(pdf_buffer.getvalue()).decode()
            href = f'<a href="data:application/pdf;base64,{b64}" download="practice_analysis_report.pdf">Download Practice Analysis Report</a>'
            st.markdown(href, unsafe_allow_html=True)

# Report generation
st.sidebar.header("Generate Report")
report_format = st.sidebar.selectbox("Choose report format", ["CSV", "PDF", "DOCX"])

if st.sidebar.button("Generate Report"):
    report_data = generate_report(G, df, analysis_type)

    if report_format == "CSV":
        # Export as CSV
        csv = report_data.to_csv(index=False) if isinstance(report_data, pd.DataFrame) else pd.DataFrame([report_data]).to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        href = f'<a href="data:file/csv;base64,{b64}" download="network_analysis_report.csv">Download CSV Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)
    
    elif report_format == "PDF":
        # Export as PDF
        pdf_buffer = io.BytesIO()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Network Analysis Report", ln=True, align='C')
        if isinstance(report_data, dict):
            for key, value in report_data.items():
                pdf.cell(200, 10, txt=f"{key}: {value}", ln=True)
        else:
            for line in report_data.to_string(index=False).split("\n"):
                pdf.cell(200, 10, txt=line, ln=True)
        pdf.output(pdf_buffer)
        pdf_buffer.seek(0)
        b64 = base64.b64encode(pdf_buffer.getvalue()).decode()
        href = f'<a href="data:application/pdf;base64,{b64}" download="network_analysis_report.pdf">Download PDF Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)

    elif report_format == "DOCX":
        # Export as DOCX
        doc_buffer = io.BytesIO()
        doc = Document()
        doc.add_heading("Network Analysis Report", 0)
        if isinstance(report_data, dict):
            for key, value in report_data.items():
                doc.add_paragraph(f"{key}: {value}")
        else:
            for line in report_data.to_string(index=False).split("\n"):
                doc.add_paragraph(line)
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        b64 = base64.b64encode(doc_buffer.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="network_analysis_report.docx">Download DOCX Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.markdown("Created by Your Nampalli Srinivas")
st.markdown("---")
st.markdown("### Note: This project is for educational purposes only.")
st.markdown("Connect with the creator:")
st.markdown("[![LinkedIn](https://img.shields.io/badge/LinkedIn-0077B5?style=for-the-badge&logo=linkedin&logoColor=white)](https://www.linkedin.com/in/srinivas-nampalli/)")
st.markdown("[![Twitter](https://img.shields.io/badge/Twitter-1DA1F2?style=for-the-badge&logo=twitter&logoColor=white)](https://x.com/Srinivas26k)")

if __name__ == "__main__":
    st.write("Streamlit app is running.")