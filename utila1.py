
import pandas as pd
import numpy as np
from sklearn.manifold import TSNE
from sklearn.preprocessing import StandardScaler
from community import community_louvain
import matplotlib.pyplot as plt
from fpdf import FPDF
from docx import Document
import random
import networkx as nx
import io
import streamlit as st
# List of proper nouns for random data generation
PROPER_NOUNS = [
    "Alice", "Bob", "Charlie", "David", "Emma", "Frank", "Grace", "Henry", "Isabel", "Jack",
    "Katherine", "Liam", "Mia", "Noah", "Olivia", "Peter", "Quinn", "Rachel", "Samuel", "Tara",
    "New York", "London", "Tokyo", "Paris", "Sydney", "Berlin", "Moscow", "Rome", "Cairo", "Delhi",
    "Amazon", "Google", "Microsoft", "Apple", "Facebook", "Tesla", "Netflix", "IBM", "Intel", "Oracle",
    "Harvard", "Stanford", "MIT", "Oxford", "Cambridge", "Princeton", "Yale", "UCLA", "Berkeley", "Columbia",
    "Python", "Java", "JavaScript", "Ruby", "C++", "Swift", "Kotlin", "PHP", "HTML", "CSS"
]


def preprocess_data(df):
    """
    Preprocess the DataFrame to identify and prepare source and target columns,
    while retaining all additional columns.
    """
    try:
        # Attempt to identify 'source' and 'target' columns
        if 'source' in df.columns and 'target' in df.columns:
            # Data is already in the correct format
            st.info("Source and target columns found.")
        elif set(df.columns) == set(['from', 'to']):
            # Rename 'from' and 'to' to 'source' and 'target'
            df = df.rename(columns={'from': 'source', 'to': 'target'})
            st.info("Renamed 'from' and 'to' columns to 'source' and 'target'.")
        elif len(df.columns) >= 2:
            # Assume the first two columns are source and target
            df.rename(columns={df.columns[0]: 'source', df.columns[1]: 'target'}, inplace=True)
            st.warning("First two columns used as 'source' and 'target'.")
        else:
            raise ValueError("Unable to determine source and target columns. Please ensure your data has clear columns.")

        # Convert to strings to ensure compatibility
        df['source'] = df['source'].astype(str)
        df['target'] = df['target'].astype(str)

        # Remove rows with missing values in 'source' or 'target'
        df.dropna(subset=['source', 'target'], inplace=True)

        st.success("Preprocessing completed successfully.")
        return df

    except Exception as e:
        raise ValueError(f"Error processing data: {str(e)}")


def handle_large_file(uploaded_file, chunksize=100000):
    """
    Handle large files by processing them in chunks.
    """
    try:
        st.info("Processing large file in chunks...")
        # Choose the appropriate reader based on file type
        if uploaded_file.name.endswith('.csv'):
            chunks = pd.read_csv(uploaded_file, chunksize=chunksize)
        elif uploaded_file.name.endswith('.xlsx'):
            raise ValueError("Excel files larger than 500 MB are not supported. Convert to CSV.")
        elif uploaded_file.name.endswith('.txt'):
            chunks = pd.read_csv(uploaded_file, sep='\t', chunksize=chunksize)
        else:
            raise ValueError("Unsupported file format. Please upload CSV or TXT.")

        # Combine chunks into a single DataFrame
        df = pd.concat(chunks)
        st.success("File loaded successfully.")
        return preprocess_data(df)

    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None


def handle_file_upload(uploaded_file):
    """
    Handle both small and large file uploads.
    """
    if uploaded_file.size > 500 * 1024 * 1024:  # 500 MB limit
        return handle_large_file(uploaded_file)
    else:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            elif uploaded_file.name.endswith('.xlsx'):
                df = pd.read_excel(uploaded_file)
            elif uploaded_file.name.endswith('.txt'):
                df = pd.read_csv(uploaded_file, sep='\t')
            else:
                st.error("Unsupported file format.")
                return None

            return preprocess_data(df)
        except Exception as e:
            st.error(f"Error processing file: {str(e)}")
            return None


def create_network(df):
    G = nx.from_pandas_edgelist(df, 'source', 'target')
    return G

def generate_random_network(num_nodes):
    G = nx.gnm_random_graph(num_nodes, num_nodes * 2)
    nodes = random.sample(PROPER_NOUNS, num_nodes)
    mapping = {i: name for i, name in enumerate(nodes)}
    G = nx.relabel_nodes(G, mapping)
    df = nx.to_pandas_edgelist(G)
    df.columns = ['source', 'target']
    return G, df

def calculate_network_metrics(G):
    metrics = {
        "Number of Nodes": G.number_of_nodes(),
        "Number of Edges": G.number_of_edges(),
        "Average Degree": sum(dict(G.degree()).values()) / G.number_of_nodes(),
        "Density": nx.density(G),
        "Average Clustering Coefficient": nx.average_clustering(G),
        "Average Shortest Path Length": nx.average_shortest_path_length(G) if nx.is_connected(G) else "N/A (Graph is not connected)"
    }
    return metrics

def find_path_between_nodes(G, source, target):
    try:
        path = nx.shortest_path(G, source=source, target=target)
        return path
    except nx.NetworkXNoPath:
        return None

def get_node_neighbors(G, node):
    return list(G.neighbors(node))

def analyze_node(G, node):
    metrics = {
        "degree": G.degree(node),
        "clustering_coefficient": nx.clustering(G, node),
        "betweenness_centrality": nx.betweenness_centrality(G)[node],
        "closeness_centrality": nx.closeness_centrality(G)[node],
        "eigenvector_centrality": nx.eigenvector_centrality(G)[node],
        "neighbors": list(G.neighbors(node))
    }
    return metrics

def get_community_structure(G):
    return community_louvain.best_partition(G)

def calculate_graph_resilience(G):
    original_size = G.number_of_nodes()
    removed_nodes = []
    largest_component_sizes = []

    while G.number_of_nodes() > 0:
        # Remove the node with the highest degree
        node_to_remove = max(G.degree, key=lambda x: x[1])[0]
        G.remove_node(node_to_remove)
        removed_nodes.append(node_to_remove)

        # Calculate the size of the largest connected component
        largest_cc = max(nx.connected_components(G), key=len)
        largest_component_sizes.append(len(largest_cc) / original_size)

    return removed_nodes, largest_component_sizes

def enhance_labels_with_gemini(G, context):
    # This is a placeholder function. In a real-world scenario, you would integrate
    # with a language model API to generate meaningful labels.
    enhanced_labels = {}
    for node in G.nodes():
        enhanced_labels[node] = f"Enhanced {node} ({context})"
    return enhanced_labels

def calculate_advanced_metrics(G):
    metrics = {
        "Assortativity Coefficient": nx.degree_assortativity_coefficient(G),
        "Graph Clique Number": nx.graph_clique_number(G),
        "Graph Number of Cliques": nx.graph_number_of_cliques(G),
        "Average Node Connectivity": nx.average_node_connectivity(G),
        "Degree Pearson Correlation Coefficient": nx.degree_pearson_correlation_coefficient(G),
    }
    return metrics

def compare_networks(G1, G2):
    metrics = {
        "nodes": (G1.number_of_nodes(), G2.number_of_nodes()),
        "edges": (G1.number_of_edges(), G2.number_of_edges()),
        "density": (nx.density(G1), nx.density(G2)),
        "avg_clustering": (nx.average_clustering(G1), nx.average_clustering(G2)),
        "avg_shortest_path": (nx.average_shortest_path_length(G1) if nx.is_connected(G1) else "N/A",
                              nx.average_shortest_path_length(G2) if nx.is_connected(G2) else "N/A"),
        "diameter": (nx.diameter(G1) if nx.is_connected(G1) else "N/A",
                     nx.diameter(G2) if nx.is_connected(G2) else "N/A"),
        "assortativity": (nx.degree_assortativity_coefficient(G1),
                          nx.degree_assortativity_coefficient(G2)),
    }
    return metrics

def dimensionality_reduction(G, method='tsne'):
    # Get the adjacency matrix
    adj_matrix = nx.to_numpy_array(G)
    
    # Normalize the adjacency matrix
    scaler = StandardScaler()
    adj_matrix_normalized = scaler.fit_transform(adj_matrix)
    
    if method.lower() == 'tsne':
        # Apply t-SNE
        tsne = TSNE(n_components=2, random_state=42)
        coords = tsne.fit_transform(adj_matrix_normalized)
    else:
        raise ValueError(f"Unsupported method: {method}")
    
    return coords

def generate_pdf_report():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Network Analysis Report", ln=True, align='C')
    pdf.cell(200, 10, txt="Generated using Streamlit", ln=True, align='C')
    
    # Add more content to the PDF as needed
    pdf.cell(200, 10, txt="This is a sample report.", ln=True, align='L')
    
    # Create a BytesIO object
    pdf_output = io.BytesIO()
    pdf.output(pdf_output)
    
    # Set the BytesIO object's position to the beginning
    pdf_output.seek(0)
    
    return pdf_output

def generate_docx_report(G, filename="network_report.docx"):
    doc = Document()
    doc.add_heading("Network Analysis Report", 0)
    
    # Add basic metrics
    metrics = calculate_network_metrics(G)
    for key, value in metrics.items():
        doc.add_paragraph(f"{key}: {value}")
    
    # Add advanced metrics
    advanced_metrics = calculate_advanced_metrics(G)
    doc.add_heading("Advanced Metrics", level=1)
    for key, value in advanced_metrics.items():
        doc.add_paragraph(f"{key}: {value}")
    
    # Save the document
    doc.save(filename)

def export_to_csv(G, filename="network_data.csv"):
    df = nx.to_pandas_edgelist(G)
    df.to_csv(filename, index=False)

def export_to_graphml(G, filename="network.graphml"):
    nx.write_graphml(G, filename)

def export_to_gml(G, filename="network.gml"):
    nx.write_gml(G, filename)