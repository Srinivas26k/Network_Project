import streamlit as st
import networkx as nx
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from pyvis.network import Network
import random
from fpdf import FPDF
from docx import Document
import io
import base64
import seaborn as sns
import matplotlib.pyplot as plt
from scipy import stats
import statsmodels.api as sm
from sklearn.preprocessing import StandardScaler
import streamlit_toggle as tog
from sklearn.manifold import TSNE
from sklearn.cluster import KMeans
import streamlit.components.v1 as components
from utila1 import (
    preprocess_data,
    create_network,
    generate_random_network,
    calculate_network_metrics,
    find_path_between_nodes,
    get_node_neighbors,
    analyze_node,
    get_community_structure,
    calculate_graph_resilience,
    enhance_labels_with_gemini,
    calculate_advanced_metrics,
    compare_networks,
    dimensionality_reduction
)

# Set page config
st.set_page_config(page_title="Network Analysis for Scientists", layout="wide")

# Initialize session state
if 'graphs' not in st.session_state:
    st.session_state.graphs = {}
if 'data' not in st.session_state:
    st.session_state.data = None
if 'labels' not in st.session_state:
    st.session_state.labels = {}



st.title("Network Analysis ")
if st.button("Support this project"):
    st.markdown("""
        <div style="position: fixed; bottom: 10px; right: 10px; background-color: white; padding: 10px; border: 1px solid #ddd; box-shadow: 0px 0px 10px rgba(0, 0, 0, 0.1); z-index: 1000;">
            <h4>Support this project</h4>
            <p>If you find this project useful, consider supporting me:</p>
            <a href="https://buymeacoffee.com/srinivaskiv" target="_blank">
                <img src="https://img.shields.io/badge/Buy%20Me%20A%20Coffee-FFDD00?style=for-the-badge&logo=buy-me-a-coffee&logoColor=black" alt="Buy Me A Coffee">
            </a>
            <h4>Connect with me</h4>
            <a href="https://www.linkedin.com/in/srinivas-nampalli/" target="_blank">
                <img src="https://img.shields.io/badge/LinkedIn-0077B5?style=for-the-badge&logo=linkedin&logoColor=white" alt="LinkedIn">
            </a>
            <a href="https://x.com/Srinivas26k" target="_blank">
                <img src="https://img.shields.io/badge/Twitter-1DA1F2?style=for-the-badge&logo=twitter&logoColor=white" alt="Twitter">
            </a>
            <h4>Report Issues</h4>
            <a href="https://github.com/Srinivas26k/Ropar_Network_Project/issues" target="_blank">
                <img src="https://img.shields.io/badge/GitHub-Issues-181717?style=for-the-badge&logo=github&logoColor=white" alt="GitHub Issues">
            </a>
        </div>
    """, unsafe_allow_html=True)
# Add the following code at the end of the file

# Add a function to explain network concepts
def explain_network_concept(concept):
    explanations = {
        "Node": "A node represents an entity in the network, such as a person, organization, or data point.",
        "Edge": "An edge represents a connection or relationship between two nodes in the network.",
        "Path": "A path is a sequence of edges that connects two nodes in the network.",
        "Trail": "A trail is a path that doesn't repeat edges but may revisit nodes.",
        "Walk": "A walk is a sequence of nodes and edges, allowing repetition of both nodes and edges.",
        "Degree": "The degree of a node is the number of edges connected to it.",
        "Centrality": "Centrality measures the importance of a node in the network based on its position.",
        "Clustering Coefficient": "The clustering coefficient measures the degree to which nodes tend to cluster together.",
        "Community": "A community is a group of nodes that are more densely connected to each other than to the rest of the network.",
        "Network Density": "Network density is the ratio of the number of edges in the network to the maximum possible number of edges.",
    }
    return explanations.get(concept, "Concept not found in the database.")

# Add a function to handle file upload and preprocessing
def handle_file_upload(uploaded_file):
    if uploaded_file.name.endswith('.csv'):
        df = pd.read_csv(uploaded_file)
    elif uploaded_file.name.endswith('.xlsx'):
        df = pd.read_excel(uploaded_file)
    elif uploaded_file.name.endswith('.txt'):
        df = pd.read_csv(uploaded_file, sep='\t')
    else:
        st.error("Unsupported file format. Please upload a CSV, Excel, or TXT file.")
        return None

    try:
        df = preprocess_data(df)
        return df
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return None

# Sidebar for data input
st.sidebar.header("Data Input")
data_source = st.sidebar.radio("Select data source:", ["Upload your own data", "Generate random data", "Compare networks"])

if data_source == "Upload your own data":
    uploaded_file = st.sidebar.file_uploader("Upload a file:", type=["csv", "xlsx", "txt"])
    if uploaded_file is not None:
        df = handle_file_upload(uploaded_file)
        if df is not None:
            G = create_network(df)
            st.session_state.graphs['main'] = G
            st.session_state.data = df
            st.session_state.labels = enhance_labels_with_gemini(G, context="uploaded data")
    else:
        st.sidebar.warning("Please upload a file to proceed.")
        st.stop()

elif data_source == "Generate random data":
    num_nodes = st.sidebar.slider("Number of Nodes", min_value=5, max_value=50, value=20)
    G, df = generate_random_network(num_nodes)
    st.session_state.graphs['main'] = G
    st.session_state.data = df
    st.session_state.labels = enhance_labels_with_gemini(G, context="random network")

elif data_source == "Compare networks":
    uploaded_file1 = st.sidebar.file_uploader("Upload first network file:", type=["csv", "xlsx", "txt"], key="file1")
    uploaded_file2 = st.sidebar.file_uploader("Upload second network file:", type=["csv", "xlsx", "txt"], key="file2")
    
    if uploaded_file1 is not None and uploaded_file2 is not None:
        df1 = handle_file_upload(uploaded_file1)
        df2 = handle_file_upload(uploaded_file2)
        if df1 is not None and df2 is not None:
            G1 = create_network(df1)
            G2 = create_network(df2)
            st.session_state.graphs['network1'] = G1
            st.session_state.graphs['network2'] = G2
    else:
        st.sidebar.warning("Please upload both files to proceed with comparison.")
        st.stop()

# Navigation
if data_source != "Compare networks":
    analysis_level = st.sidebar.radio("Choose Analysis Level", ["Basic Analysis", "Advanced Analysis"])

    if analysis_level == "Basic Analysis":
        analysis_type = st.sidebar.selectbox(
            "Choose Analysis Type",
            ["Overview", "Node Analysis", "Path Analysis", "Degree Distribution", "Network Visualization"]
        )
    else:
        analysis_type = st.sidebar.selectbox(
            "Choose Analysis Type",
            ["Statistical Analysis", "Advanced Visualizations", "Network Comparison", "Scientific Metrics", "Research Tools"]
        )
else:
    analysis_type = "Network Comparison"

# Add descriptions and explanations for each analysis type
analysis_descriptions = {
    "Overview": "Provides basic network statistics and a visual representation of the network structure.",
    "Node Analysis": "Examines individual nodes and their relationships within the network.",
    "Path Analysis": "Analyzes the paths between nodes and calculates related metrics.",
    "Degree Distribution": "Visualizes and analyzes the distribution of node degrees in the network.",
    "Network Visualization": "Offers various layouts to visualize the network structure.",
    "Statistical Analysis": "Performs advanced statistical analysis on network properties.",
    "Advanced Visualizations": "Provides sophisticated visualization techniques for network analysis.",
    "Network Comparison": "Compares two networks based on various metrics and properties.",
    "Scientific Metrics": "Calculates advanced metrics used in scientific network analysis.",
    "Research Tools": "Offers tools for generating reports, exporting data, and performing advanced analysis."
}

st.sidebar.markdown("---")
st.sidebar.subheader("Analysis Description")
st.sidebar.write(analysis_descriptions.get(analysis_type, ""))

# Add a section for network concept explanations
st.sidebar.markdown("---")
concepts = {
    "Node": "A node represents an entity in the network, such as a person, organization, or data point.",
    "Edge": "An edge represents a connection or relationship between two nodes in the network.",
    "Path": "A path is a sequence of edges that connects two nodes in the network.",
    "Trail": "A trail is a path that doesn't repeat edges but may revisit nodes.",
    "Walk": "A walk is a sequence of nodes and edges, allowing repetition of both nodes and edges.",
    "Degree": "The degree of a node is the number of edges connected to it.",
    "Centrality": "Centrality measures the importance of a node in the network based on its position.",
    "Clustering Coefficient": "The clustering coefficient measures the degree to which nodes tend to cluster together.",
    "Community": "A community is a group of nodes that are more densely connected to each other than to the rest of the network.",
    "Network Density": "Network density is the ratio of the number of edges in the network to the maximum possible number of edges.",
}

selected_concept = st.sidebar.selectbox("Select a concept to learn more:", list(concepts.keys()))
st.sidebar.write(concepts.get(selected_concept, "Concept not found in the database."))

st.sidebar.write(explain_network_concept(selected_concept))

# Main content area
st.title("Network Analysis")
st.write(f"Current Analysis: {analysis_type}")
st.write(analysis_descriptions.get(analysis_type, ""))

# Implement error handling for each analysis type
try:
    if analysis_type == "Overview":
        st.write("Basic network statistics and visualization")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Network Statistics")
                metrics = calculate_network_metrics(G)
                for k, v in metrics.items():
                    st.write(f"{k}: {v}")
            
            with col2:
                st.subheader("Network Visualization")
                net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
                
                # Add nodes with proper string conversion
                for node in G.nodes():
                    net.add_node(str(node), label=str(node), title=str(node))
                
                # Add edges with proper string conversion
                for edge in G.edges():
                    net.add_edge(str(edge[0]), str(edge[1]))
                
                net.save_graph("network.html")
                HtmlFile = open("network.html", 'r', encoding='utf-8')
                source_code = HtmlFile.read()
                components.html(source_code, height=400)

    elif analysis_type == "Node Analysis":
        st.write("Analyze specific nodes and their relationships")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            # Node search
            search_node = st.selectbox("Select a node to analyze:", list(G.nodes()))
            
            if search_node:
                col1, col2 = st.columns(2)
                
                with col1:
                    st.subheader("Node Metrics")
                    node_metrics = analyze_node(G, search_node)
                    for metric, value in node_metrics.items():
                        if metric != "neighbors":
                            st.write(f"{metric.replace('_', ' ').title()}: {value}")
                    
                    st.subheader("Neighbors")
                    st.write(", ".join(map(str, node_metrics["neighbors"])))
                
                with col2:
                    st.subheader("Local Network Visualization")
                    local_net = Network(height="300px", width="100%", bgcolor="#222222", font_color="white")
                    
                    # Add central node and its neighbors
                    local_net.add_node(str(search_node), color='red', size=20, title=str(search_node))
                    for neighbor in node_metrics["neighbors"]:
                        local_net.add_node(str(neighbor), size=15, title=str(neighbor))
                        local_net.add_edge(str(search_node), str(neighbor))
                    
                    local_net.save_graph("local_network.html")
                    HtmlFile = open("local_network.html", 'r', encoding='utf-8')
                    source_code = HtmlFile.read()
                    components.html(source_code, height=300)

    elif analysis_type == "Path Analysis":
        st.write("Analyze paths between nodes")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            col1, col2 = st.columns(2)
            with col1:
                source = st.selectbox("Source node:", list(G.nodes()), key="source")
            with col2:
                target = st.selectbox("Target node:", list(G.nodes()), key="target")
            
            if st.button("Find Path"):
                path = find_path_between_nodes(G, source, target)
                if path:
                    st.write("Shortest path:", " → ".join(map(str, path)))
                    
                    # Visualize path
                    path_net = Network(height="300px", width="100%", bgcolor="#222222", font_color="white")
                    
                    # Add nodes and edges along the path
                    for i in range(len(path)):
                        path_net.add_node(str(path[i]), color='lightblue', title=str(path[i]))
                        if i > 0:
                            path_net.add_edge(str(path[i-1]), str(path[i]))
                    
                    path_net.save_graph("path.html")
                    HtmlFile = open("path.html", 'r', encoding='utf-8')
                    source_code = HtmlFile.read()
                    components.html(source_code, height=300)
                    
                    # Calculate path probability
                    if len(path) > 1:
                        path_prob = 1 / (len(path) - 1)
                        st.write(f"Path Probability: {path_prob:.4f}")
                    else:
                        st.write("Path Probability: N/A (source and target are the same node)")
                    
                    # Calculate number of trails
                    num_trails = nx.number_of_simple_paths(G, source, target)
                    st.write(f"Number of Simple Paths: {num_trails}")
                else:
                    st.warning("No path found between selected nodes.")

    elif analysis_type == "Degree Distribution":
        st.write("Analyze the degree distribution of the network")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            degrees = [d for n, d in G.degree()]
            
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
            
            # Histogram
            ax1.hist(degrees, bins='auto', alpha=0.7)
            ax1.set_title("Degree Distribution")
            ax1.set_xlabel("Degree")
            ax1.set_ylabel("Frequency")
            
            # Cumulative Distribution
            sorted_degrees = sorted(degrees, reverse=True)
            cumulative = np.arange(1, len(sorted_degrees) + 1) / len(sorted_degrees)
            ax2.plot(sorted_degrees, cumulative, 'b-')
            ax2.set_title("Cumulative Degree Distribution")
            ax2.set_xlabel("Degree")
            ax2.set_ylabel("Cumulative Frequency")
            ax2.set_xscale('log')
            ax2.set_yscale('log')
            
            st.pyplot(fig)
            
            # Basic statistics
            st.write(f"Average Degree: {np.mean(degrees):.2f}")
            st.write(f"Median Degree: {np.median(degrees):.2f}")
            st.write(f"Maximum Degree: {np.max(degrees)}")
            st.write(f"Minimum Degree: {np.min(degrees)}")

    elif analysis_type == "Network Visualization":
        st.write("Visualize the network with different layouts")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            layout_option = st.selectbox("Choose Layout", ["Spring", "Circular", "Random", "Shell"])
            
            fig, ax = plt.subplots(figsize=(10, 8))
            
            if layout_option == "Spring":
                pos = nx.spring_layout(G)
            elif layout_option == "Circular":
                pos = nx.circular_layout(G)
            elif layout_option == "Random":
                pos = nx.random_layout(G)
            else:  # Shell
                pos = nx.shell_layout(G)
            
            nx.draw(G, pos, with_labels=True, node_color='lightblue', node_size=500, font_size=8, font_weight='bold', ax=ax)
            ax.set_title(f"{layout_option} Layout")
            st.pyplot(fig)
            
            # Additional information
            st.write(f"Number of Nodes: {G.number_of_nodes()}")
            st.write(f"Number of Edges: {G.number_of_edges()}")
            st.write(f"Network Density: {nx.density(G):.4f}")

    elif analysis_type == "Statistical Analysis":
        st.write("Advanced statistical analysis of network properties")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            analysis_option = st.selectbox(
                "Choose Analysis Type",
                ["Degree Distribution", "Clustering Analysis", "Centrality Analysis", "Statistical Tests"]
            )
            
            if analysis_option == "Degree Distribution":
                degrees = [d for n, d in G.degree()]
                
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
                
                # KDE Plot
                sns.kdeplot(data=degrees, ax=ax1)
                ax1.set_title("Degree Distribution (KDE)")
                
                # ECDF Plot
                sns.ecdfplot(data=degrees, ax=ax2)
                ax2.set_title("Empirical Cumulative Distribution")
                
                st.pyplot(fig)
                
                # Statistical summary
                st.write("Statistical Summary:")
                summary = pd.DataFrame({
                    'Statistic': ['Mean', 'Median', 'Std', 'Skewness', 'Kurtosis'],
                    'Value': [
                        np.mean(degrees),
                        np.median(degrees),
                        np.std(degrees),
                        stats.skew(degrees),
                        stats.kurtosis(degrees)
                    ]
                })
                st.write(summary)

            elif analysis_option == "Clustering Analysis":
                clustering_coeffs = list(nx.clustering(G).values())
                
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
                
                # Histogram with KDE
                sns.histplot(data=clustering_coeffs, kde=True, ax=ax1)
                ax1.set_title("Clustering Coefficient Distribution")
                
                # Box plot
                sns.boxplot(data=clustering_coeffs, ax=ax2)
                ax2.set_title("Clustering Coefficient Box Plot")
                
                st.pyplot(fig)

            elif analysis_option == "Centrality Analysis":
                degree_cent = nx.degree_centrality(G)
                betweenness_cent = nx.betweenness_centrality(G)
                closeness_cent = nx.closeness_centrality(G)
                
                centrality_df = pd.DataFrame({
                    'Node': list(G.nodes()),
                    'Degree': list(degree_cent.values()),
                    'Betweenness': list(betweenness_cent.values()),
                    'Closeness': list(closeness_cent.values())
                })
                
                fig = px.scatter_matrix(centrality_df, 
                                      dimensions=['Degree', 'Betweenness', 'Closeness'],
                                      title="Centrality Measures Correlation")
                st.plotly_chart(fig)

            elif analysis_option == "Statistical Tests":
                st.subheader("Statistical Tests")
                
                # Normality test
                degrees = [d for n, d in G.degree()]
                stat, p_value = stats.normaltest(degrees)
                st.write(f"Normality Test (D'Agostino and Pearson's test):")
                st.write(f"Statistic: {stat:.4f}")
                st.write(f"p-value: {p_value:.4f}")
                
                # QQ Plot
                fig, ax = plt.subplots(figsize=(8, 6))
                stats.probplot(degrees, dist="norm", plot=ax)
                ax.set_title("Q-Q Plot of Degree Distribution")
                st.pyplot(fig)

    elif analysis_type == "Advanced Visualizations":
        st.write("Advanced network visualization techniques")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            viz_option = st.selectbox(
                "Choose Visualization",
                ["Network Layout Comparison", "3D Network Plot", "Heat Map", "Advanced Statistical Plots", "Network Embedding Visualization"]
            )
            
            if viz_option == "Network Layout Comparison":
                layouts = ["spring", "circular", "random", "shell"]
                cols = st.columns(2)
                for i, layout in enumerate(layouts):
                    with cols[i % 2]:
                        st.write(f"{layout.capitalize()} Layout")
                        net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
                        pos = getattr(nx, f"{layout}_layout")(G)
                        
                        # Add nodes with positions
                        for node, position in pos.items():
                            net.add_node(str(node), label=str(node), x=float(position[0])*1000, y=float(position[1])*1000, title=str(node))
                        
                        # Add edges
                        for edge in G.edges():
                            net.add_edge(str(edge[0]), str(edge[1]))
                        
                        net.save_graph(f"{layout}_layout.html")
                        HtmlFile = open(f"{layout}_layout.html", 'r', encoding='utf-8')
                        source_code = HtmlFile.read()
                        components.html(source_code, height=400)
                        
                        # Save as PNG
                        plt.figure(figsize=(10, 8))
                        nx.draw(G, pos, with_labels=True, node_color='lightblue', node_size=500, font_size=8, font_weight='bold')
                        plt.title(f"{layout.capitalize()} Layout")
                        plt.savefig(f"{layout}_layout.png")
                        st.image(f"{layout}_layout.png")

            elif viz_option == "3D Network Plot":
                pos_3d = nx.spring_layout(G, dim=3)
                
                # Create 3D scatter plot
                edge_x = []
                edge_y = []
                edge_z = []
                for edge in G.edges():
                    x0, y0, z0 = pos_3d[edge[0]]
                    x1, y1, z1 = pos_3d[edge[1]]
                    edge_x.extend([x0, x1, None])
                    edge_y.extend([y0, y1, None])
                    edge_z.extend([z0, z1, None])

                # Create the 3D network plot
                fig = go.Figure(data=[
                    go.Scatter3d(x=edge_x, y=edge_y, z=edge_z,
                                mode='lines',
                                line=dict(color='#888', width=1),
                                hoverinfo='none'
                    ),
                    go.Scatter3d(
                        x=[pos_3d[node][0] for node in G.nodes()],
                        y=[pos_3d[node][1] for node in G.nodes()],
                        z=[pos_3d[node][2] for node in G.nodes()],
                        mode='markers+text',
                        text=list(G.nodes()),
                        marker=dict(size=6, color='lightblue')
                    )
                ])
                
                fig.update_layout(showlegend=False, width=800, height=800)
                st.plotly_chart(fig)

            elif viz_option == "Heat Map":
                # Create adjacency matrix
                adj_matrix = nx.adjacency_matrix(G).todense()
                
                fig, ax = plt.subplots(figsize=(10, 8))
                sns.heatmap(adj_matrix, cmap='YlOrRd', ax=ax)
                ax.set_title("Network Adjacency Matrix Heatmap")
                st.pyplot(fig)

            elif viz_option == "Advanced Statistical Plots":
                st.subheader("Advanced Statistical Visualizations")
                
                plot_type = st.selectbox(
                    "Select Plot Type",
                    ["Joint Plot", "Violin Plot", "Strip Plot", "Swarm Plot"]
                )
                
                # Prepare data
                node_data = pd.DataFrame({
                    'Degree': [d for n, d in G.degree()],
                    'Clustering': list(nx.clustering(G).values()),
                    'Centrality': list(nx.degree_centrality(G).values())
                })
                
                if plot_type == "Joint Plot":
                    g = sns.jointplot(data=node_data, x="Degree", y="Clustering",
                                    kind="reg", height=8)
                    st.pyplot(g.fig)
                
                elif plot_type == "Violin Plot":
                    fig, ax = plt.subplots(figsize=(10, 6))
                    sns.violinplot(data=node_data, ax=ax)
                    ax.set_title("Distribution of Network Metrics")
                    st.pyplot(fig)
                
                elif plot_type == "Strip Plot":
                    fig, ax = plt.subplots(figsize=(10, 6))
                    sns.stripplot(data=node_data, ax=ax)
                    ax.set_title("Individual Values of Network Metrics")
                    st.pyplot(fig)
                
                elif plot_type == "Swarm Plot":
                    fig, ax = plt.subplots(figsize=(10, 6))
                    sns.swarmplot(data=node_data, ax=ax)
                    ax.set_title("Distributed Strip Plot of Network Metrics")
                    st.pyplot(fig)

            elif viz_option == "Network Embedding Visualization":
                st.subheader("Network Embedding Visualization")
                
                method = st.selectbox("Choose embedding method", ["t-sne"])
                
                try:
                    coords = dimensionality_reduction(G, method.lower())
                    
                    fig = px.scatter(x=coords[:, 0], y=coords[:, 1], 
                                     hover_name=list(G.nodes()),
                                     title=f"Network Embedding using {method.upper()}")
                    st.plotly_chart(fig)

                    # Optional: Add clustering
                    if st.checkbox("Show clusters"):
                        n_clusters = st.slider("Number of clusters", 2, 10, 5)
                        kmeans = KMeans(n_clusters=n_clusters, random_state=42)
                        clusters = kmeans.fit_predict(coords)
                        
                        fig = px.scatter(x=coords[:, 0], y=coords[:, 1], 
                                         color=clusters,
                                         hover_name=list(G.nodes()),
                                         title=f"Network Embedding using {method.upper()} with {n_clusters} clusters")
                        st.plotly_chart(fig)
                except Exception as e:
                    st.error(f"Error in embedding visualization: {str(e)}")

    elif analysis_type == "Network Comparison":
        st.write("Compare two networks")
        
        if 'network1' in st.session_state.graphs and 'network2' in st.session_state.graphs:
            G1 = st.session_state.graphs['network1']
            G2 = st.session_state.graphs['network2']
            
            comparison_type = st.selectbox(
                "Choose Comparison Type",
                ["Basic Metrics", "Degree Distribution", "Centrality Comparison", "Statistical Tests", "Advanced Comparison"]
            )
            
            if comparison_type == "Basic Metrics":
                try:
                    metrics = compare_networks(G1, G2)
                    
                    metrics_df = pd.DataFrame({
                        'Metric': list(metrics.keys()),
                        'Network 1': [m[0] for m in metrics.values()],
                        'Network 2': [m[1] for m in metrics.values()]
                    })
                    
                    st.write(metrics_df)
                    
                    # Ensure metrics are numeric before plotting
                    metrics_df = metrics_df.apply(pd.to_numeric, errors='coerce')
                    
                    # Visualize the comparison
                    fig, ax = plt.subplots(figsize=(10, 6))
                    metrics_df.plot(x='Metric', y=['Network 1', 'Network 2'], kind='bar', ax=ax)
                    plt.xticks(rotation=45)
                    st.pyplot(fig)
                    st.write("This chart compares basic network metrics between the two uploaded networks.")
                except Exception as e:
                    st.error(f"Error in network comparison: {str(e)}")

            elif comparison_type == "Degree Distribution":
                try:
                    degrees1 = [d for n, d in G1.degree()]
                    degrees2 = [d for n, d in G2.degree()]
                    
                    # Ensure both degree lists are of the same length
                    max_len = max(len(degrees1), len(degrees2))
                    degrees1.extend([0] * (max_len - len(degrees1)))
                    degrees2.extend([0] * (max_len - len(degrees2)))
                    
                    if len(degrees1) == 0 or len(degrees2) == 0:
                        st.warning("One or both of the networks have no nodes.")
                    else:
                        fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
                        
                        sns.kdeplot(data=degrees1, ax=ax1, label='Network 1')
                        sns.kdeplot(data=degrees2, ax=ax1, label='Network 2')
                        ax1.set_title("Degree Distribution Comparison (KDE)")
                        ax1.legend()
                        
                        plot_data = pd.DataFrame({
                            'Network 1': degrees1,
                            'Network 2': degrees2
                        })
                        sns.boxplot(data=plot_data, ax=ax2)
                        ax2.set_title("Degree Distribution Comparison (Box Plot)")
                        
                        st.pyplot(fig)
                        st.write("These plots show the degree distribution comparison between the two networks using Kernel Density Estimation (KDE) and box plots.")
                        
                        stat, pval = stats.ks_2samp(degrees1, degrees2)
                        st.write("Kolmogorov-Smirnov test for degree distributions:")
                        st.write(f"Statistic: {stat:.4f}")
                        st.write(f"p-value: {pval:.4f}")
                        st.write("The Kolmogorov-Smirnov test compares the degree distributions of the two networks. A small p-value suggests that the distributions are significantly different.")
                except Exception as e:
                    st.error(f"Error in degree distribution comparison: {str(e)}")

            elif comparison_type == "Centrality Comparison":
                try:
                    metrics1 = {
                        'Degree': nx.degree_centrality(G1),
                        'Betweenness': nx.betweenness_centrality(G1),
                        'Closeness': nx.closeness_centrality(G1)
                    }
                    
                    metrics2 = {
                        'Degree': nx.degree_centrality(G2),
                        'Betweenness': nx.betweenness_centrality(G2),
                        'Closeness': nx.closeness_centrality(G2)
                    }
                    
                    # Ensure both centrality metrics are of the same length
                    max_len = max(len(metrics1['Degree']), len(metrics2['Degree']))
                    for metric in metrics1.keys():
                        metrics1[metric] = list(metrics1[metric].values()) + [0] * (max_len - len(metrics1[metric]))
                        metrics2[metric] = list(metrics2[metric].values()) + [0] * (max_len - len(metrics2[metric]))
                    
                    fig, axes = plt.subplots(1, 3, figsize=(15, 5))
                    
                    for i, metric in enumerate(metrics1.keys()):
                        data = pd.DataFrame({
                            'Network 1': metrics1[metric],
                            'Network 2': metrics2[metric]
                        })
                        sns.violinplot(data=data, ax=axes[i])
                        axes[i].set_title(f"{metric} Centrality")
                    
                    plt.tight_layout()
                    st.pyplot(fig)
                    st.write("These violin plots compare the distribution of centrality measures (Degree, Betweenness, and Closeness) between the two networks.")
                except Exception as e:
                    st.error(f"Error in centrality comparison: {str(e)}")

            elif comparison_type == "Statistical Tests":
                st.subheader("Statistical Comparison Tests")
                
                degrees1 = [d for n, d in G1.degree()]
                degrees2 = [d for n, d in G2.degree()]
                
                stat, pval = stats.mannwhitneyu(degrees1, degrees2)
                st.write("Mann-Whitney U test for degree distributions:")
                st.write(f"Statistic: {stat:.4f}")
                st.write(f"p-value: {pval:.4f}")
                st.write("The Mann-Whitney U test compares the degree distributions of the two networks. A small p-value suggests that the distributions are significantly different.")
                
                d = (np.mean(degrees1) - np.mean(degrees2)) / np.sqrt((np.std(degrees1)**2 + np.std(degrees2)**2) / 2)
                st.write(f"Cohen's d effect size: {d:.4f}")
                st.write("Cohen's d measures the standardized difference between the means of the two degree distributions. Values around 0.2, 0.5, and 0.8 represent small, medium, and large effect sizes, respectively.")
                
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
                stats.probplot(degrees1, dist="norm", plot=ax1)
                ax1.set_title("Q-Q Plot (Network 1)")
                stats.probplot(degrees2, dist="norm", plot=ax2)
                ax2.set_title("Q-Q Plot (Network 2)")
                st.pyplot(fig)
                st.write("These Q-Q plots compare the degree distributions of each network to a normal distribution. Deviations from the diagonal line indicate departures from normality.")
            
            elif comparison_type == "Advanced Comparison":
                st.subheader("Advanced Network Comparison")
                
                # Community structure comparison
                communities1 = list(nx.community.greedy_modularity_communities(G1))
                communities2 = list(nx.community.greedy_modularity_communities(G2))
                
                modularity1 = nx.community.modularity(G1, communities1)
                modularity2 = nx.community.modularity(G2, communities2)
                
                st.write(f"Network 1 - Number of Communities: {len(communities1)}, Modularity: {modularity1:.4f}")
                st.write(f"Network 2 - Number of Communities: {len(communities2)}, Modularity: {modularity2:.4f}")
                st.write("These metrics compare the community structure of the two networks. Higher modularity values indicate more distinct community structures.")
                
                # Spectral properties comparison
                L1 = nx.normalized_laplacian_matrix(G1)
                L2 = nx.normalized_laplacian_matrix(G2)
                
                eigenvalues1 = np.linalg.eigvals(L1.toarray())
                eigenvalues2 = np.linalg.eigvals(L2.toarray())
                
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
                ax1.hist(eigenvalues1.real, bins='auto', alpha=0.7, label='Network 1')
                ax1.hist(eigenvalues2.real, bins='auto', alpha=0.7, label='Network 2')
                ax1.set_title("Eigenvalue Distribution Comparison")
                ax1.set_xlabel("Eigenvalue")
                ax1.set_ylabel("Frequency")
                ax1.legend()
                
                ax2.plot(range(len(eigenvalues1)), sorted(eigenvalues1.real), label='Network 1')
                ax2.plot(range(len(eigenvalues2)), sorted(eigenvalues2.real), label='Network 2')
                ax2.set_title("Eigenvalue Spectrum Comparison")
                ax2.set_xlabel("Index")
                ax2.set_ylabel("Eigenvalue")
                ax2.legend()
                st.pyplot(fig)
                st.write("These plots compare the spectral properties of the two networks. The eigenvalue distribution and spectrum can reveal differences in network structure and connectivity patterns.")
                
                # Network resilience comparison
                def calculate_largest_component_size(g):
                    return len(max(nx.connected_components(g), key=len))
                
                def simulate_node_removal(G, n_removals):
                    G_copy = G.copy()
                    sizes = []
                    for _ in range(n_removals):
                        sizes.append(calculate_largest_component_size(G_copy))
                        if len(G_copy.nodes()) > 0:
                            G_copy.remove_node(random.choice(list(G_copy.nodes())))
                    return sizes
                
                n_removals = min(G1.number_of_nodes(), G2.number_of_nodes(), 20)
                resilience1 = simulate_node_removal(G1, n_removals)
                resilience2 = simulate_node_removal(G2, n_removals)
                
                fig, ax = plt.subplots(figsize=(10, 6))
                ax.plot(range(n_removals), resilience1, label='Network 1')
                ax.plot(range(n_removals), resilience2, label='Network 2')
                ax.set_xlabel('Number of Nodes Removed')
                ax.set_ylabel('Size of Largest Component')
                ax.set_title('Network Resilience Comparison')
                ax.legend()
                st.pyplot(fig)
                st.write("This plot compares the resilience of the two networks to random node removal. Networks that maintain larger connected components after node removal are considered more resilient.")

    elif analysis_type == "Scientific Metrics":
        st.write("Advanced metrics for scientific analysis")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            metric_type = st.selectbox(
                "Choose Metric Type",
                ["Topological Analysis", "Spectral Analysis", "Community Detection", "Network Resilience"]
            )
            
            if metric_type == "Topological Analysis":
                st.subheader("Topological Metrics")
                
                # Calculate metrics
                metrics = {
                    "Average Path Length": nx.average_shortest_path_length(G) if nx.is_connected(G) else "N/A",
                    "Graph Diameter": nx.diameter(G) if nx.is_connected(G) else "N/A",
                    "Graph Radius": nx.radius(G) if nx.is_connected(G) else "N/A",
                    "Average Clustering Coefficient": nx.average_clustering(G),
                    "Graph Density": nx.density(G),
                    "Graph Transitivity": nx.transitivity(G)
                }
                
                for k, v in metrics.items():
                    st.write(f"{k}: {v}")
                
                # Visualize degree distribution on log-log scale
                degrees = [d for n, d in G.degree()]
                fig, ax = plt.subplots(figsize=(10, 6))
                plt.hist(degrees, bins='auto', density=True, alpha=0.7)
                plt.xscale('log')
                plt.yscale('log')
                plt.title("Degree Distribution (Log-Log Scale)")
                plt.xlabel("Degree")
                plt.ylabel("Frequency")
                st.pyplot(fig)

            elif metric_type == "Spectral Analysis":
                st.subheader("Spectral Analysis")
                
                # Calculate eigenvalues
                L = nx.normalized_laplacian_matrix(G)
                eigenvalues = np.linalg.eigvals(L.toarray())
                
                # Plot eigenvalue distribution
                fig, ax = plt.subplots(figsize=(10, 6))
                plt.hist(eigenvalues.real, bins='auto', density=True)
                plt.title("Eigenvalue Distribution of Normalized Laplacian")
                plt.xlabel("Eigenvalue")
                plt.ylabel("Density")
                st.pyplot(fig)
                
                # Spectral gap
                sorted_eigenvals = np.sort(eigenvalues.real)
                spectral_gap = sorted_eigenvals[1] - sorted_eigenvals[0]
                st.write(f"Spectral Gap: {spectral_gap:.4f}")

            elif metric_type == "Community Detection":
                st.subheader("CommunityCommunity Detection Analysis")
                
                # Detect communities
                communities = list(nx.community.greedy_modularity_communities(G))
                
                # Calculate modularity
                modularity = nx.community.modularity(G, communities)
                st.write(f"Number of Communities: {len(communities)}")
                st.write(f"Modularity Score: {modularity:.4f}")
                
                # Visualize communities
                pos = nx.spring_layout(G)
                fig, ax = plt.subplots(figsize=(10, 8))
                
                colors = plt.cm.rainbow(np.linspace(0, 1, len(communities)))
                for comm, color in zip(communities, colors):
                    nx.draw_networkx_nodes(G, pos, comm, node_color=[color], ax=ax)
                
                nx.draw_networkx_edges(G, pos, alpha=0.5, ax=ax)
                plt.title("Community Structure Visualization")
                st.pyplot(fig)

            elif metric_type == "Network Resilience":
                st.subheader("Network Resilience Analysis")
                
                # Calculate node removal impact
                def calculate_largest_component_size(g):
                    return len(max(nx.connected_components(g), key=len))
                
                # Random node removal
                random_removal = []
                G_random = G.copy()
                nodes = list(G_random.nodes())
                np.random.shuffle(nodes)
                
                for i in range(min(20, len(nodes))):
                    size = calculate_largest_component_size(G_random)
                    random_removal.append(size)
                    if nodes:
                        G_random.remove_node(nodes.pop())
                
                # Targeted node removal (highest degree first)
                targeted_removal = []
                G_targeted = G.copy()
                nodes = sorted(G_targeted.degree(), key=lambda x: x[1], reverse=True)
                
                for i in range(min(20, len(nodes))):
                    size = calculate_largest_component_size(G_targeted)
                    targeted_removal.append(size)
                    if nodes:
                        G_targeted.remove_node(nodes[0][0])
                        nodes = sorted(G_targeted.degree(), key=lambda x: x[1], reverse=True)
                
                # Plot results
                fig, ax = plt.subplots(figsize=(10, 6))
                plt.plot(range(len(random_removal)), random_removal, label='Random Removal')
                plt.plot(range(len(targeted_removal)), targeted_removal, label='Targeted Removal')
                plt.xlabel('Number of Nodes Removed')
                plt.ylabel('Size of Largest Component')
                plt.title('Network Resilience Analysis')
                plt.legend()
                st.pyplot(fig)

    elif analysis_type == "Research Tools":
        st.write("Tools for research and analysis")
    
        if 'main' in st.session_state.graphs:
            G = st.session_state.graphs['main']
            
            tool_type = st.selectbox(
                "Choose Research Tool",
                ["Network Analysis Report", "Data Export", "Statistical Summary", "Advanced Metrics"]
            )
            
            if tool_type == "Network Analysis Report":
                report_format = st.selectbox("Choose report format", ["PDF", "DOCX"])
                
                if st.button("Generate Report"):
                    if report_format == "PDF":
                        pdf_buffer = io.BytesIO()
                        pdf = FPDF()
                        pdf.add_page()
                        pdf.set_font("Arial", size=12)
                        
                        # Add content to PDF
                        pdf.cell(200, 10, txt="Network Analysis Report", ln=True, align='C')
                        metrics = calculate_network_metrics(G)
                        for key, value in metrics.items():
                            pdf.cell(200, 10, txt=f"{key}: {value}", ln=True)
                        
                        pdf_buffer.write(pdf.output(dest='S').encode('latin1'))
                        pdf_buffer.seek(0)
                        
                        # Create download button
                        b64_pdf = base64.b64encode(pdf_buffer.getvalue()).decode()
                        href = f'<a href="data:application/pdf;base64,{b64_pdf}" download="network_analysis_report.pdf">Download PDF Report</a>'
                        st.markdown(href, unsafe_allow_html=True)
                    
                    else:  # DOCX
                        doc_buffer = io.BytesIO()
                        doc = Document()
                        doc.add_heading("Network Analysis Report", 0)
                        
                        # Add content to DOCX
                        metrics = calculate_network_metrics(G)
                        for key, value in metrics.items():
                            doc.add_paragraph(f"{key}: {value}")
                        
                        doc.save(doc_buffer)
                        doc_buffer.seek(0)
                        doc_buffer.seek(0)
                        
                        # Create download button
                        b64_doc = base64.b64encode(doc_buffer.getvalue()).decode()
                        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64_doc}" download="network_analysis_report.docx">Download DOCX Report</a>'
                        st.markdown(href, unsafe_allow_html=True)

            elif tool_type == "Data Export":
                export_format = st.selectbox("Choose export format", ["CSV", "GraphML", "GML"])
                
                if st.button("Export Data"):
                    if export_format == "CSV":
                        edges_df = nx.to_pandas_edgelist(G)
                        csv = edges_df.to_csv(index=False)
                        b64 = base64.b64encode(csv.encode()).decode()
                        href = f'<a href="data:file/csv;base64,{b64}" download="network_data.csv">Download CSV</a>'
                        st.markdown(href, unsafe_allow_html=True)
                    
                    elif export_format == "GraphML":
                        buffer = io.StringIO()
                        nx.write_graphml(G, buffer)
                        b64 = base64.b64encode(buffer.getvalue().encode()).decode()
                        href = f'<a href="data:text/graphml;base64,{b64}" download="network.graphml">Download GraphML</a>'
                        st.markdown(href, unsafe_allow_html=True)
                    
                    else:  # GML
                        buffer = io.StringIO()
                        nx.write_gml(G, buffer)
                        b64 = base64.b64encode(buffer.getvalue().encode()).decode()
                        href = f'<a href="data:text/gml;base64,{b64}" download="network.gml">Download GML</a>'
                        st.markdown(href, unsafe_allow_html=True)

            elif tool_type == "Statistical Summary":
                st.subheader("Comprehensive Statistical Summary")
                
                # Node-level statistics
                degrees = [d for n, d in G.degree()]
                clustering_coeffs = list(nx.clustering(G).values())
                
                stats_df = pd.DataFrame({
                    'Metric': ['Degree', 'Clustering Coefficient'],
                    'Mean': [np.mean(degrees), np.mean(clustering_coeffs)],
                    'Median': [np.median(degrees), np.median(clustering_coeffs)],
                    'Std Dev': [np.std(degrees), np.std(clustering_coeffs)],
                    'Min': [np.min(degrees), np.min(clustering_coeffs)],
                    'Max': [np.max(degrees), np.max(clustering_coeffs)]
                })
                
                st.write(stats_df)
                
                # Correlation analysis
                st.subheader("Correlation Analysis")
                degree_cent = nx.degree_centrality(G)
                between_cent = nx.betweenness_centrality(G)
                close_cent = nx.closeness_centrality(G)
                
                corr_df = pd.DataFrame({
                    'Degree': list(degree_cent.values()),
                    'Betweenness': list(between_cent.values()),
                    'Closeness': list(close_cent.values())
                })
                
                fig, ax = plt.subplots(figsize=(8, 6))
                sns.heatmap(corr_df.corr(), annot=True, cmap='coolwarm', ax=ax)
                plt.title("Correlation Matrix of Centrality Measures")
                st.pyplot(fig)

            elif tool_type == "Advanced Metrics":
                st.subheader("Advanced Network Metrics")
                
                # Calculate advanced metrics
                metrics = calculate_advanced_metrics(G)
                
                # Display metrics
                for metric, value in metrics.items():
                    st.write(f"{metric}: {value:.4f}")
                
                # Advanced visualizations
                st.subheader("Advanced Metric Visualizations")
                
                # Eigenvector analysis
                L = nx.normalized_laplacian_matrix(G)
                eigenvalues = np.linalg.eigvals(L.toarray())
                
                fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 5))
                
                # Eigenvalue spectrum
                ax1.plot(range(len(eigenvalues)), sorted(eigenvalues.real), 'b.')
                ax1.set_title("Eigenvalue Spectrum")
                ax1.set_xlabel("Index")
                ax1.set_ylabel("Eigenvalue")
                
                # Scree plot
                explained_var = np.sort(np.abs(eigenvalues))[::-1]
                explained_var_ratio = explained_var / explained_var.sum()
                ax2.plot(range(1, len(explained_var_ratio) + 1), np.cumsum(explained_var_ratio))
                ax2.set_title("Scree Plot")
                ax2.set_xlabel("Number of Components")
                ax2.set_ylabel("Cumulative Explained Variance Ratio")
                
                plt.tight_layout()
                st.pyplot(fig)

except Exception as e:
    st.error(f"An error occurred during the analysis: {str(e)}")
    st.write("Please try again with different parameters or contact support if the issue persists.")

# Contribution and social links
# Contribution and social links
st.markdown("---")
st.markdown("Created by Nampalli Srinivas")
st.markdown("---")
st.markdown("### Support this project")
st.markdown("[![Buy Me A Coffee](https://img.shields.io/badge/Buy%20Me%20A%20Coffee-FFDD00?style=for-the-badge&logo=buy-me-a-coffee&logoColor=black)](https://buymeacoffee.com/srinivaskiv)")
st.markdown("---")
st.markdown("### Connect with me")
st.markdown("[![LinkedIn](https://img.shields.io/badge/LinkedIn-0077B5?style=for-the-badge&logo=linkedin&logoColor=white)](https://www.linkedin.com/in/srinivas-nampalli/)")
st.markdown("[![Twitter](https://img.shields.io/badge/Twitter-1DA1F2?style=for-the-badge&logo=twitter&logoColor=white)](https://x.com/Srinivas26k)")
st.markdown("### Report Issues")
st.markdown("[![GitHub Issues](https://img.shields.io/badge/GitHub-Issues-181717?style=for-the-badge&logo=github&logoColor=white)](https://github.com/Srinivas26k/Ropar_Network_Project/issues)")
st.markdown("---")
if __name__ == "__main__":
    st.write("""
I extend my heartfelt gratitude to the esteemed [Sudarshan Iyengar Sir](https://www.linkedin.com/in/sudarshan-iyengar-3560b8145/) for teaching me and offering a unique perspective on AI.  
A special thanks to my friends [Prakhar Gupta](https://www.linkedin.com/in/prakhar-kselis/), [Jinal Gupta](https://www.linkedin.com/in/jinal-gupta-220a652b6/), and Purba Bandyopadhyay for constantly motivating and encouraging me to take on such a wonderful project.  
Your guidance and support have been truly invaluable!
""")


    st.markdown("---")
    st.write("Network Analysis app is running.")

