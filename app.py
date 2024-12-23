import streamlit as st
import networkx as nx
import pandas as pd
import plotly.graph_objects as go
from pyvis.network import Network
import faiss
import numpy as np
from fpdf import FPDF
from docx import Document
import io
import base64

from utils import preprocess_data, create_network, create_embeddings, query_faiss, generate_report, calculate_pagerank, analyze_connection_quality, generate_insights

# Set up FAISS
dimension = 1  # Dimension of the simple embeddings
index = faiss.IndexFlatL2(dimension)

st.set_page_config(page_title="Interactive Network Analysis", layout="wide")
st.title("Interactive Network Analysis with AI-Powered Insights")

# Load and preprocess data
@st.cache_data
def load_data():
    # Assuming the CSV file is in the same directory as the script
    df = pd.read_csv("impression_network.csv")
    return preprocess_data(df)

df = load_data()


# Create network graph
G = create_network(df)

# Create embeddings and add to FAISS index
embeddings = create_embeddings(df['Name'].tolist())
index.add(embeddings)

# Sidebar for navigation
st.sidebar.header("Navigation")
analysis_type = st.sidebar.selectbox(
    "Choose Analysis Type",
    ["Overview", "Concept Explanation", "Individual Node Analysis", "Influence Paths", "Community Detection", "Connection Quality", "RAG Insights", "Gamification"]
)

# Main content area
st.header(analysis_type)

if analysis_type == "Overview":
    st.write("This section provides a basic overview of the network, including statistics and visualization.")
    
    # Display basic network statistics
    st.subheader("Network Statistics")
    st.write(f"Number of nodes: {G.number_of_nodes()}")
    st.write(f"Number of edges: {G.number_of_edges()}")
    st.write(f"Average degree: {2 * G.number_of_edges() / G.number_of_nodes():.2f}")
    
    # Create and display network visualization
    st.subheader("Network Visualization")
    net = Network(height="600px", width="100%", bgcolor="#222222", font_color="white")
    net.from_nx(G)
    net.save_graph("network.html")
    st.components.v1.html(open("network.html", "r").read(), height=600)

elif analysis_type == "Concept Explanation":
    st.write("Learn about key concepts in network analysis.")
    
    concepts = {
        "Node": "A fundamental unit in a network, representing an entity (e.g., a person, organization, or data point).",
        "Edge": "A connection between two nodes, representing a relationship or interaction.",
        "Degree": "The number of edges connected to a node, indicating its connectivity.",
        "Path": "A sequence of nodes and edges connecting two nodes in the network.",
        "Centrality": "A measure of a node's importance in the network, based on its position and connections.",
        "Community": "A group of nodes that are more densely connected to each other than to the rest of the network.",
        "PageRank": "An algorithm that measures the importance of nodes based on the structure of incoming links."
    }
    
    for concept, explanation in concepts.items():
        st.subheader(concept)
        st.write(explanation)

elif analysis_type == "Individual Node Analysis":
    st.write("Analyze individual nodes in the network to understand their properties and importance.")
    
    selected_node = st.selectbox("Select a node to analyze", df['Name'].tolist())
    
    st.subheader("Node Properties")
    degree = G.degree(selected_node)
    st.write(f"Degree: {degree}")
    
    component = next(nx.connected_components(G))
    if selected_node in component:
        st.write(f"Connected component size: {len(component)}")
    else:
        st.write("The selected node is not part of the largest connected component.")
    
    pagerank = calculate_pagerank(G)
    st.write(f"PageRank score: {pagerank[selected_node]:.6f}")
    
    st.subheader("Node Neighbors")
    neighbors = list(G.neighbors(selected_node))
    st.write(f"Neighbors: {', '.join(neighbors)}")

elif analysis_type == "Influence Paths":
    st.write("Discover the shortest paths between nodes to understand influence flow in the network.")
    
    source = st.selectbox("Select source node", df['Name'].tolist())
    target = st.selectbox("Select target node", df['Name'].tolist())
    
    if st.button("Find Influence Path"):
        try:
            path = nx.shortest_path(G, source=source, target=target)
            st.success(f"Influence path from {source} to {target}:")
            st.write(" -> ".join(path))
            
            # Visualize the path
            path_graph = G.subgraph(path)
            net = Network(height="400px", width="100%", bgcolor="#222222", font_color="white")
            net.from_nx(path_graph)
            net.save_graph("path.html")
            st.components.v1.html(open("path.html", "r").read(), height=400)
        except nx.NetworkXNoPath:
            st.error(f"No path found between {source} and {target}")

elif analysis_type == "Community Detection":
    st.write("Identify and visualize communities within the network.")
    
    communities = nx.community.louvain_communities(G)
    
    st.subheader("Community Statistics")
    st.write(f"Number of communities detected: {len(communities)}")
    
    # Visualize communities
    st.subheader("Community Visualization")
    pos = nx.spring_layout(G)
    fig = go.Figure()
    
    for i, community in enumerate(communities):
        node_x = [pos[node][0] for node in community]
        node_y = [pos[node][1] for node in community]
        fig.add_trace(go.Scatter(x=node_x, y=node_y, mode='markers', name=f'Community {i+1}',
                                 marker=dict(size=10, line=dict(width=2))))
    
    fig.update_layout(title="Community Detection Visualization", showlegend=True)
    st.plotly_chart(fig)

elif analysis_type == "Connection Quality":
    st.write("Evaluate the quality of connections for a specific node in the network.")
    
    selected_node = st.selectbox("Select a node to analyze", df['Name'].tolist())
    quality_score, explanation = analyze_connection_quality(G, selected_node)
    
    st.subheader("Connection Quality Analysis")
    st.write(f"Connection Quality Score: {quality_score:.2f}")
    st.write(f"Explanation: {explanation}")
    
    st.subheader("Node Connections")
    neighbors = list(G.neighbors(selected_node))
    st.write(f"Neighbors of {selected_node}: {', '.join(neighbors)}")
    
    st.subheader("Connections between Neighbors")
    for i, n1 in enumerate(neighbors):
        for n2 in neighbors[i+1:]:
            if G.has_edge(n1, n2):
                st.write(f"- {n1} is connected to {n2}")

elif analysis_type == "RAG Insights":
    st.write("Get AI-powered insights about the network using Retrieval Augmented Generation (RAG).")
    
    query = st.text_input("Ask a question about the network:")
    
    if query:
        relevant_nodes = query_faiss(index, query, df['Name'].tolist(), k=5)
        st.subheader("Relevant Nodes")
        st.write(", ".join(relevant_nodes))
        
        st.subheader("Network Insights")
        insights = generate_insights(G, query, relevant_nodes)
        st.write(insights)

elif analysis_type == "Gamification":
    st.write("Experiment with the network by adding or removing nodes and observe the impact.")
    
    action = st.radio("Choose an action:", ["Add Node", "Remove Node"])
    
    if action == "Add Node":
        new_node = st.text_input("Enter the name of the new node:")
        connect_to = st.multiselect("Connect to existing nodes:", df['Name'].tolist())
        
        if st.button("Add Node"):
            G.add_node(new_node)
            for node in connect_to:
                G.add_edge(new_node, node)
            st.success(f"Added node {new_node} with {len(connect_to)} connections.")
    else:
        remove_node = st.selectbox("Select a node to remove:", df['Name'].tolist())
        
        if st.button("Remove Node"):
            G.remove_node(remove_node)
            st.success(f"Removed node {remove_node}.")
    
    st.subheader("Updated Network Statistics")
    st.write(f"Number of nodes: {G.number_of_nodes()}")
    st.write(f"Number of edges: {G.number_of_edges()}")
    
    st.subheader("Updated PageRank Scores")
    updated_pagerank = calculate_pagerank(G)
    for node, score in sorted(updated_pagerank.items(), key=lambda x: x[1], reverse=True)[:5]:
        st.write(f"{node}: {score:.6f}")

# Report generation
st.sidebar.header("Generate Report")
report_format = st.sidebar.selectbox("Choose report format", ["CSV", "PDF", "DOCX"])

if st.sidebar.button("Generate Report"):
    report_data = generate_report(G, df, analysis_type)
    
    if report_format == "CSV":
        csv = report_data.to_csv(index=False)
        b64 = base64.b64encode(csv.encode()).decode()
        href = f'<a href="data:file/csv;base64,{b64}" download="network_analysis_report.csv">Download CSV Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)
    elif report_format == "PDF":
        pdf_buffer = io.BytesIO()
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Network Analysis Report", ln=True, align='C')
        for key, value in report_data.items():
            pdf.cell(200, 10, txt=f"{key}: {value}", ln=True)
        pdf.output(pdf_buffer)
        b64 = base64.b64encode(pdf_buffer.getvalue()).decode()
        href = f'<a href="data:application/pdf;base64,{b64}" download="network_analysis_report.pdf">Download PDF Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)
    elif report_format == "DOCX":
        doc_buffer = io.BytesIO()
        doc = Document()
        doc.add_heading("Network Analysis Report", 0)
        for key, value in report_data.items():
            doc.add_paragraph(f"{key}: {value}")
        doc.save(doc_buffer)
        b64 = base64.b64encode(doc_buffer.getvalue()).decode()
        href = f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="network_analysis_report.docx">Download DOCX Report</a>'
        st.sidebar.markdown(href, unsafe_allow_html=True)

st.sidebar.markdown("---")
st.sidebar.markdown("Created by Your Name")
st.markdown("---")
st.markdown("### Note: This project is for educational purposes only.")
st.markdown("Connect with the creator:")
st.markdown("[![LinkedIn](https://img.shields.io/badge/LinkedIn-0077B5?style=for-the-badge&logo=linkedin&logoColor=white)](https://www.linkedin.com/in/srinivas-nampalli/)")
st.markdown("[![Twitter](https://img.shields.io/badge/Twitter-1DA1F2?style=for-the-badge&logo=twitter&logoColor=white)](https://x.com/Srinivas26k)")

if __name__ == "__main__":
    st.write("Streamlit app is running.")

